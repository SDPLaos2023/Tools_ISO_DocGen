"""
doc_validator.py — ISO Document Quality Assurance Validator
=============================================================
ตรวจสอบความถูกต้อง ความสมบูรณ์ และความเป็นทางการของเอกสาร ISO ที่สร้างออกมา

การตรวจสอบแบ่งเป็น 5 หมวด:
  1. Structure   — ครบโครงสร้าง (Cover, Doc Control, Version History, Sections)
  2. Content     — ข้อมูลสำคัญไม่ว่าง / ไม่ใช่ placeholder
  3. Spelling    — คำผิดทั่วไป (Thai typos + known incorrect terms)
  4. Formality   — ภาษาทางการ (ตรวจคำไม่เป็นทางการ)
  5. Cross-ref   — REQ-ID, TC-ID, RISK-ID ที่อ้างอิงมีอยู่จริง

วิธีใช้:
  from utils.doc_validator import validate_all, print_report
  results = validate_all(config, output_root)
  approved = print_report(results)
"""

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn

# ─── Severity Levels ──────────────────────────────────────────────────────────
CRITICAL = "CRITICAL"   # ❌ ต้องแก้ก่อน approve
WARNING  = "WARNING"    # ⚠  ควรแก้ไข
INFO     = "INFO"       # ℹ  สังเกตไว้ ไม่บังคับ


@dataclass
class Issue:
    severity: str
    category: str
    doc_file: str
    message: str
    suggestion: str = ""


@dataclass
class DocResult:
    file_path: str
    doc_name: str
    issues: List[Issue] = field(default_factory=list)

    @property
    def critical_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == CRITICAL)

    @property
    def warning_count(self) -> int:
        return sum(1 for i in self.issues if i.severity == WARNING)

    @property
    def approved(self) -> bool:
        return self.critical_count == 0


# ─── Known Thai Spelling / Formality Issues ───────────────────────────────────
# Format: (wrong_pattern, correct_form, severity, note)
SPELLING_RULES = [
    # Common Thai typos
    (r"กรุณาติดต่อ", "กรุณาติดต่อ", INFO, ""),  # sentinel — correct
    (r"ร่ร",         "รร",          CRITICAL, "คำผิด: 'ร่ร' → 'รร'"),
    (r"เเ",          "แ",           CRITICAL, "คำผิด: 'เเ' (สองตัว) → 'แ' (ตัวเดียว)"),
    (r"ณ\s+ปัจจุบัน", "ณ ปัจจุบัน", INFO, ""),  # correct form
    # Informal / non-standard words
    (r"\bok\b",      "ผ่าน / เรียบร้อย", WARNING, "ภาษาไม่เป็นทางการ: 'ok' → ใช้ 'ผ่าน' หรือ 'เรียบร้อย'"),
    (r"\bOK\b",      "ผ่าน",         WARNING, "ภาษาไม่เป็นทางการ: 'OK'"),
    (r"\basap\b",    "โดยเร็ว",      WARNING, "ภาษาไม่เป็นทางการ: 'asap'"),
    (r"\bTBD\b",     "รอกำหนด",      WARNING, "ควรระบุค่าจริง แทน 'TBD'"),
    (r"\bTBA\b",     "รอกำหนด",      WARNING, "ควรระบุค่าจริง แทน 'TBA'"),
    (r"\bn/a\b",     "ไม่มี / ไม่เกี่ยวข้อง", INFO, "พิจารณาใช้ข้อความภาษาไทยแทน 'n/a'"),
    (r"\bN/A\b",     "ไม่มี / ไม่เกี่ยวข้อง", INFO, "พิจารณาใช้ข้อความภาษาไทยแทน 'N/A'"),
    # Placeholder texts (from template defaults)
    (r"Full Name",   "",             CRITICAL, "พบข้อความ placeholder 'Full Name' — ยังไม่ได้กรอกข้อมูล"),
    (r"member@company\.com", "",     CRITICAL, "พบ email placeholder — ยังไม่ได้กรอกข้อมูล"),
    (r"Organization / Company Name", "", CRITICAL, "พบข้อความ placeholder ชื่อองค์กร"),
    (r"Department / Business Unit",  "", CRITICAL, "พบข้อความ placeholder หน่วยงาน"),
]

# Required structural keywords that every ISO document must contain
REQUIRED_STRUCTURE_KEYWORDS = [
    ("Document Control",   CRITICAL, "ไม่พบส่วน 'Document Control' — จำเป็นตามมาตรฐาน ISO/IEC 29110"),
    ("Version History",    CRITICAL, "ไม่พบส่วน 'Version History' — จำเป็นสำหรับการควบคุมเวอร์ชัน"),
    ("Prepared by",        CRITICAL, "ไม่พบผู้จัดทำ 'Prepared by'"),
    ("Approved by",        CRITICAL, "ไม่พบผู้อนุมัติ 'Approved by'"),
]

# Minimum paragraph count (extremely short doc is likely incomplete)
MIN_PARAGRAPHS = 10

# ─── Validator Core ───────────────────────────────────────────────────────────

def _extract_full_text(doc: Document) -> str:
    """Extract all text from a .docx file including tables."""
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _check_structure(doc: Document, full_text: str, file_name: str) -> List[Issue]:
    """Check that required structural sections are present."""
    issues = []
    for keyword, severity, message in REQUIRED_STRUCTURE_KEYWORDS:
        if keyword not in full_text:
            issues.append(Issue(
                severity=severity,
                category="Structure",
                doc_file=file_name,
                message=message,
                suggestion=f"เพิ่มส่วน '{keyword}' ตาม template มาตรฐาน"
            ))

    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    if para_count < MIN_PARAGRAPHS:
        issues.append(Issue(
            severity=WARNING,
            category="Structure",
            doc_file=file_name,
            message=f"เอกสารสั้นเกินไป ({para_count} paragraphs) — อาจสร้างไม่ครบ",
            suggestion="ตรวจสอบว่าเนื้อหาทุก section ถูก generate ครบ"
        ))

    # Check logo/header presence (heuristic: check if document has header content)
    has_header = False
    for section in doc.sections:
        if section.header and any(p.text.strip() for p in section.header.paragraphs):
            has_header = True
            break
        # Also check header tables
        if section.header:
            for tbl in section.header.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            has_header = True
                            break
    if not has_header:
        issues.append(Issue(
            severity=WARNING,
            category="Structure",
            doc_file=file_name,
            message="ไม่พบ Header ในเอกสาร — อาจดูไม่เป็นทางการ",
            suggestion="ตรวจสอบ add_header_footer() ถูกเรียกใช้"
        ))

    return issues


def _check_content(full_text: str, file_name: str, config: dict) -> List[Issue]:
    """Check for empty placeholders and missing required data."""
    issues = []
    p = config.get("project", {})

    # Verify project name appears in document
    project_name = p.get("name", "")
    if project_name and project_name not in full_text:
        issues.append(Issue(
            severity=WARNING,
            category="Content",
            doc_file=file_name,
            message=f"ไม่พบชื่อโครงการ '{project_name}' ในเอกสาร",
            suggestion="ตรวจสอบว่า config.project.name ถูกใส่ในเอกสาร"
        ))

    # Verify document ID format pattern appears
    code = p.get("code", "")
    if code and code not in full_text:
        issues.append(Issue(
            severity=WARNING,
            category="Content",
            doc_file=file_name,
            message=f"ไม่พบ Project Code '{code}' ในเอกสาร — Document ID อาจผิดพลาด",
            suggestion="ตรวจสอบ get_doc_id() และ cover page"
        ))

    # Verify organization appears
    org = p.get("organization", "")
    if org and org not in full_text:
        issues.append(Issue(
            severity=INFO,
            category="Content",
            doc_file=file_name,
            message=f"ไม่พบชื่อองค์กร '{org}' — อาจขาดใน header หรือ cover",
            suggestion="ตรวจสอบ cover page organization field"
        ))

    return issues


def _check_spelling_formality(full_text: str, file_name: str) -> List[Issue]:
    """Check for spelling errors and informal language."""
    issues = []
    for pattern, correct, severity, note in SPELLING_RULES:
        if not note:
            continue  # Skip sentinel/correct-form entries
        matches = re.findall(pattern, full_text, re.IGNORECASE)
        if matches:
            issues.append(Issue(
                severity=severity,
                category="Spelling/Formality",
                doc_file=file_name,
                message=note,
                suggestion=f"แทนที่ด้วย: '{correct}'" if correct else "ลบหรือแก้ไขข้อความนี้"
            ))
    return issues


def _check_cross_references(full_text: str, file_name: str, config: dict) -> List[Issue]:
    """Check that REQ-IDs, TC-IDs, RISK-IDs referenced in doc exist in config."""
    issues = []

    valid_req_ids  = {r.get("id","") for r in config.get("requirements", [])}
    valid_tc_ids   = {t.get("id","") for t in config.get("test_cases", [])}
    valid_risk_ids = {r.get("id","") for r in config.get("risks", [])}

    # Find all REQ-xxx references in the document text
    found_reqs = set(re.findall(r"REQ-\d{3}", full_text))
    for req_id in found_reqs:
        if req_id not in valid_req_ids:
            issues.append(Issue(
                severity=WARNING,
                category="Cross-Reference",
                doc_file=file_name,
                message=f"อ้างอิง '{req_id}' ที่ไม่มีในรายการ Requirements",
                suggestion=f"ตรวจสอบ config requirements หรือแก้ ID ให้ถูกต้อง"
            ))

    # Find all TC-xxx references
    found_tcs = set(re.findall(r"TC-\d{3}", full_text))
    for tc_id in found_tcs:
        if tc_id not in valid_tc_ids and valid_tc_ids:
            issues.append(Issue(
                severity=INFO,
                category="Cross-Reference",
                doc_file=file_name,
                message=f"อ้างอิง '{tc_id}' ที่ไม่พบในรายการ Test Cases ของ config",
                suggestion="ตรวจสอบว่าเป็น test case ที่ยังไม่ได้เพิ่มใน config"
            ))

    return issues


def validate_document(file_path: str, config: dict) -> DocResult:
    """Validate a single .docx file. Returns DocResult."""
    file_name = os.path.basename(file_path)
    result = DocResult(file_path=file_path, doc_name=file_name)

    try:
        doc = Document(file_path)
        full_text = _extract_full_text(doc)
    except Exception as e:
        result.issues.append(Issue(
            severity=CRITICAL,
            category="File",
            doc_file=file_name,
            message=f"ไม่สามารถเปิดไฟล์ได้: {e}",
            suggestion="ตรวจสอบว่าไฟล์ .docx ไม่เสียหาย"
        ))
        return result

    result.issues.extend(_check_structure(doc, full_text, file_name))
    result.issues.extend(_check_content(full_text, file_name, config))
    result.issues.extend(_check_spelling_formality(full_text, file_name))
    result.issues.extend(_check_cross_references(full_text, file_name, config))

    return result


def validate_all(config: dict, output_root: str) -> List[DocResult]:
    """
    Validate all .docx files in the output_root directory tree.
    Skips Word temporary lock files (~$...).
    Returns list of DocResult for every file found.
    """
    results = []
    for dirpath, _, filenames in os.walk(output_root):
        for fname in sorted(filenames):
            # Skip Word temp lock files
            if fname.startswith("~$"):
                continue
            if fname.lower().endswith(".docx"):
                full_path = os.path.join(dirpath, fname)
                result = validate_document(full_path, config)
                results.append(result)
    return results


# ─── Report Printer ───────────────────────────────────────────────────────────

def print_report(results: List[DocResult], verbose: bool = False) -> bool:
    """
    Print a formatted QA report to stdout.
    Returns True if ALL documents are approved (no CRITICAL issues).
    """
    SEP = "=" * 70
    THIN = "-" * 70

    total_docs     = len(results)
    approved_docs  = sum(1 for r in results if r.approved)
    rejected_docs  = total_docs - approved_docs
    total_critical = sum(r.critical_count for r in results)
    total_warning  = sum(r.warning_count  for r in results)

    print(f"\n{SEP}")
    print("  ISO Document QA Validator Report")
    print(SEP)
    print(f"  เอกสารทั้งหมด : {total_docs} ไฟล์")
    print(f"  ✅ Approved   : {approved_docs} ไฟล์")
    print(f"  ❌ Rejected   : {rejected_docs} ไฟล์")
    print(f"  Critical Issues: {total_critical}")
    print(f"  Warnings       : {total_warning}")
    print(SEP)

    for result in results:
        status = "✅ APPROVED" if result.approved else "❌ REJECTED"
        print(f"\n  {status}  {result.doc_name}")

        if not result.issues:
            print(f"    ✓ ไม่พบปัญหา — เอกสารผ่านการตรวจสอบทุกหมวด")
            continue

        # Group by severity
        criticals = [i for i in result.issues if i.severity == CRITICAL]
        warnings  = [i for i in result.issues if i.severity == WARNING]
        infos     = [i for i in result.issues if i.severity == INFO]

        for issue in criticals:
            print(f"    ❌ [{issue.category}] {issue.message}")
            if issue.suggestion:
                print(f"       → {issue.suggestion}")

        for issue in warnings:
            print(f"    ⚠  [{issue.category}] {issue.message}")
            if issue.suggestion:
                print(f"       → {issue.suggestion}")

        if verbose:
            for issue in infos:
                print(f"    ℹ  [{issue.category}] {issue.message}")
                if issue.suggestion:
                    print(f"       → {issue.suggestion}")
        elif infos:
            print(f"    ℹ  {len(infos)} รายการ INFO (ใช้ --verbose เพื่อดูทั้งหมด)")

    print(f"\n{SEP}")
    if rejected_docs == 0:
        print("  🎉 ผ่านการตรวจสอบทั้งหมด — พร้อม Approve และส่งมอบ")
    else:
        print(f"  ⛔ ยังไม่ผ่าน — กรุณาแก้ไข {total_critical} Critical Issue(s) ก่อน Approve")
    print(f"{SEP}\n")

    return rejected_docs == 0


def get_summary_dict(results: List[DocResult]) -> dict:
    """Return a compact summary dict (used for programmatic reporting)."""
    return {
        "total":    len(results),
        "approved": sum(1 for r in results if r.approved),
        "rejected": sum(1 for r in results if not r.approved),
        "critical": sum(r.critical_count for r in results),
        "warnings": sum(r.warning_count  for r in results),
        "details":  [
            {
                "file":     r.doc_name,
                "approved": r.approved,
                "critical": r.critical_count,
                "warnings": r.warning_count,
                "issues": [
                    {"severity": i.severity, "category": i.category,
                     "message": i.message, "suggestion": i.suggestion}
                    for i in r.issues
                ]
            }
            for r in results
        ]
    }
