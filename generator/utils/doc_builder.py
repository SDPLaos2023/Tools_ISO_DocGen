"""
doc_builder.py — Common helper functions for building ISO .docx documents
All templates import from here to ensure consistent styling across all documents.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os
from datetime import datetime

# ─── Logo Path (auto-detect from project root) ────────────────────────────────
_THIS_DIR   = os.path.dirname(os.path.abspath(__file__))
_GENERATOR  = os.path.dirname(_THIS_DIR)           # generator/
_PROJECT_ROOT = os.path.dirname(_GENERATOR)        # DocISOGen/
LOGO_PATH = os.path.join(_PROJECT_ROOT, "Logo", "LogoCompany.png")


# ─── Color Palette — ISO Formatting Standard ─────────────────────────────────
# Ref: ISO Document Formatting Standard v1.0
COLOR_PRIMARY    = RGBColor(0x00, 0x33, 0x66)   # #003366 Dark Navy (headings)
COLOR_SECONDARY  = RGBColor(0xD9, 0xD9, 0xD9)   # #D9D9D9 Light Gray (table headers)
COLOR_ACCENT     = RGBColor(0xFF, 0xC0, 0x00)   # #FFC000 Gold (cover accent bar)
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)   # #FFFFFF
COLOR_LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)   # #F2F2F2 Alt row / code bg
COLOR_DARK_GRAY  = RGBColor(0x40, 0x40, 0x40)   # #404040 Secondary text
COLOR_BLACK      = RGBColor(0x00, 0x00, 0x00)   # #000000 Body text
COLOR_RED        = RGBColor(0xC0, 0x00, 0x00)   # #C00000 Error / Critical
COLOR_GREEN      = RGBColor(0x37, 0x86, 0x10)   # #378610 OK / Pass
COLOR_LINK       = RGBColor(0x00, 0x66, 0xCC)   # #0066CC Hyperlinks

# ─── Font Constants ───────────────────────────────────────────────────────────
FONT_BODY   = "Cambria"       # Primary font (ISO standard)
FONT_THAI   = "TH Sarabun New" # Thai fallback (complex script)
FONT_CODE   = "Courier New"   # Monospace for code/technical data

# ─── Font Size Constants (pt) ─────────────────────────────────────────────────
SZ_TITLE    = 20   # Cover page title
SZ_H1       = 16   # Heading level 1
SZ_H2       = 14   # Heading level 2
SZ_H3       = 12   # Heading level 3
SZ_BODY     = 11   # Body text
SZ_TABLE    = 11   # Table cell text
SZ_CAPTION  = 10   # Table/figure captions (Italic)
SZ_FOOTNOTE = 10   # Footnotes
SZ_CODE     = 10   # Code blocks
SZ_HEADER   = 10   # Page header/footer


def new_document() -> Document:
    """
    Create a new Document with ISO-standard page setup.
    A4, margins: Left 30mm | Right 20mm | Top 25mm | Bottom 25mm
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(3.0)    # 30 mm — wider for binding
    section.right_margin  = Cm(2.0)    # 20 mm
    section.top_margin    = Cm(2.5)    # 25 mm
    section.bottom_margin = Cm(2.5)    # 25 mm
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    _set_default_font(doc)
    _set_paragraph_spacing(doc)
    return doc


def _set_paragraph_spacing(doc: Document):
    """
    ISO standard paragraph spacing:
    space_before=0pt, space_after=8pt, line_spacing=1.15x (≈ 14pt at 11pt body)
    """
    style = doc.styles["Normal"]
    pf = style.paragraph_format
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pf.space_before = Pt(0)
    pf.space_after  = Pt(8)
    # 1.15× line spacing via XML (lineRule=auto, line=276 = 240*1.15)
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"),     "276")   # 240 * 1.15 = 276 twips
    spacing.set(qn("w:lineRule"), "auto")


def add_header_footer(doc: Document, config: dict, doc_id: str, doc_title: str):
    """
    Corporate-style header and footer.
    Header: [Logo] | Doc Title (bold navy) / Org (gray) | Doc ID (navy)
            ── thin gold rule ──
    Footer: ── thin navy rule ──
            Classification | Page X / Y | © Year  Org
    """
    p_cfg = config["project"]
    org   = p_cfg.get("organization", "")
    clf   = p_cfg.get("classification", "Internal")
    date  = p_cfg.get("document_date", datetime.now().strftime("%Y-%m-%d"))

    section = doc.sections[0]

    # Enable "Different First Page" — cover page gets no header/footer
    section.different_first_page_header_footer = True
    # First-page header: leave blank (cover page handles its own design)
    fph = section.first_page_header
    for para in fph.paragraphs:
        para.clear()
    # First-page footer: leave blank
    fpf = section.first_page_footer
    for para in fpf.paragraphs:
        para.clear()

    # ── HEADER (page 2 onwards) ──────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    for para in header.paragraphs:
        para.clear()

    # Logo row  (3 cells: logo | title+org | doc_id)
    htbl = header.add_table(rows=1, cols=3, width=Cm(16))
    htbl.style = "Table Grid"
    _remove_table_borders(htbl)
    htbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    htbl.rows[0].cells[0].width = Cm(2.0)
    htbl.rows[0].cells[1].width = Cm(10.0)
    htbl.rows[0].cells[2].width = Cm(4.0)

    left_cell, mid_cell, right_cell = htbl.rows[0].cells

    # Left — Logo (small)
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after  = Pt(2)
    if os.path.exists(LOGO_PATH):
        r = lp.add_run()
        r.add_picture(LOGO_PATH, height=Cm(0.8))
    else:
        r = lp.add_run(org[:2].upper())
        _apply_font(r, size=10, bold=True, color=COLOR_PRIMARY)

    # Center — Title (bold navy) + Org (small gray)
    mp = mid_cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after  = Pt(0)
    t_run = mp.add_run(doc_title.replace("\n", " "))
    _apply_font(t_run, size=10, bold=True, color=COLOR_PRIMARY)
    mp.add_run("\n")
    o_run = mp.add_run(org)
    _apply_font(o_run, size=8, color=COLOR_DARK_GRAY)

    # Right — Doc ID (right-aligned, navy)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    id_run = rp.add_run(doc_id)
    _apply_font(id_run, size=9, color=COLOR_PRIMARY)

    # Gold rule under header
    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after  = Pt(0)
    _add_para_bottom_border(rule_p, color_hex=_rgb_hex(COLOR_ACCENT))

    # ── FOOTER ──────────────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        para.clear()

    # Navy rule above footer
    rule_f = footer.add_paragraph()
    rule_f.paragraph_format.space_before = Pt(0)
    rule_f.paragraph_format.space_after  = Pt(2)
    _add_para_top_border(rule_f, color_hex=_rgb_hex(COLOR_PRIMARY))

    # 3-cell footer row
    ftbl = footer.add_table(rows=1, cols=3, width=Cm(16))
    ftbl.style = "Table Grid"
    _remove_table_borders(ftbl)
    ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    ftbl.rows[0].cells[0].width = Cm(4.0)
    ftbl.rows[0].cells[1].width = Cm(8.0)
    ftbl.rows[0].cells[2].width = Cm(4.0)

    fl_cell, fm_cell, fr_cell = ftbl.rows[0].cells

    # Left — classification label (plain text, no emoji for Word compatibility)
    flp = fl_cell.paragraphs[0]
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flp.paragraph_format.space_before = Pt(2)
    clf_badge = flp.add_run(f"[{clf}]")
    _apply_font(clf_badge, size=9, bold=True, color=COLOR_PRIMARY)

    # Center — Page X / Y
    fmp = fm_cell.paragraphs[0]
    fmp.paragraph_format.space_before = Pt(2)
    _add_page_number(fmp)

    # Right — © Year  Org
    frp = fr_cell.paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frp.paragraph_format.space_before = Pt(2)
    year = date[:4] if date else datetime.now().strftime("%Y")
    fr_run = frp.add_run(f"© {year}  {org}")
    _apply_font(fr_run, size=9, color=COLOR_DARK_GRAY)


def _set_default_font(doc: Document):
    """
    Set default body font to Cambria 11pt (ISO standard).
    Also sets TH Sarabun New as the complex-script (Thai) font via XML
    so Thai characters render correctly without manual override.
    """
    style = doc.styles["Normal"]
    font  = style.font
    font.name = FONT_BODY
    font.size = Pt(SZ_BODY)
    font.color.rgb = COLOR_BLACK

    # Set complex-script font for Thai via rFonts XML
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),     FONT_BODY)
    rFonts.set(qn("w:hAnsi"),     FONT_BODY)
    rFonts.set(qn("w:cs"),        FONT_THAI)  # Thai complex script
    rFonts.set(qn("w:eastAsia"),  FONT_THAI)


# ─── Cover Page ───────────────────────────────────────────────────────────────
def add_cover_page(doc: Document, config: dict, doc_id: str, doc_title: str,
                   doc_subtitle: str = ""):
    """
    Corporate-style cover page:
    ┌──────────────────────────────────────────────┐
    │  [Logo]   Org name           (header strip)  │
    ├──────────────────────────────────────────────┤
    │                                              │
    │  [GOLD ACCENT] Project Code                  │
    │                                              │
    │  DOCUMENT TITLE  (large)                     │
    │  subtitle         (smaller)                  │
    │                                              │
    │  ┌─────────────────────────┐                 │
    │  │ Label      │ Value      │                 │
    │  │ ...        │ ...        │                 │
    │  └─────────────────────────┘                 │
    │                                              │
    │  ISO/IEC 29110 note  (footer strip)          │
    └──────────────────────────────────────────────┘
    """
    p = config["project"]

    # ══════════════════════════════════════════════════════════════════
    # 1. TOP STRIP — Full-width Navy bar with Logo + Org name
    # ══════════════════════════════════════════════════════════════════
    top_tbl = doc.add_table(rows=1, cols=2)
    top_tbl.style = "Table Grid"
    _remove_table_borders(top_tbl)

    # Set full-page width (A4 content width ≈ 16cm after margins)
    top_tbl_width = Cm(16)

    logo_col = top_tbl.rows[0].cells[0]
    org_col  = top_tbl.rows[0].cells[1]
    logo_col.width = Cm(3.0)
    org_col.width  = Cm(13.0)

    # Navy background on both cells
    _set_cell_bg(logo_col, COLOR_PRIMARY)
    _set_cell_bg(org_col,  COLOR_PRIMARY)

    # Logo in left cell
    lp = logo_col.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after  = Pt(6)
    if os.path.exists(LOGO_PATH):
        r = lp.add_run()
        r.add_picture(LOGO_PATH, height=Cm(1.2))
    else:
        r = lp.add_run(p.get("organization","")[:2].upper())
        _apply_font(r, size=16, bold=True, color=COLOR_WHITE)

    # Org + Dept in right cell
    op = org_col.paragraphs[0]
    op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    op.paragraph_format.space_before = Pt(6)
    op.paragraph_format.space_after  = Pt(2)
    org_run = op.add_run(p.get("organization", ""))
    _apply_font(org_run, size=14, bold=True, color=COLOR_WHITE)

    op2 = org_col.add_paragraph()
    op2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    op2.paragraph_format.space_before = Pt(0)
    op2.paragraph_format.space_after  = Pt(6)
    dept_run = op2.add_run(p.get("department", ""))
    _apply_font(dept_run, size=10, color=RGBColor(0xBF, 0xD7, 0xFF))  # light blue-white

    # ══════════════════════════════════════════════════════════════════
    # 2. GOLD ACCENT BAR
    # ══════════════════════════════════════════════════════════════════
    _add_color_bar(doc, COLOR_ACCENT, height_cm=0.25)

    # ══════════════════════════════════════════════════════════════════
    # 3. PROJECT CODE TAG (เฉพาะ project code + project name — ไม่ซ้ำ doc_id)
    # ══════════════════════════════════════════════════════════════════
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tag_p.paragraph_format.space_before = Pt(20)
    tag_p.paragraph_format.space_after  = Pt(4)
    tag_run = tag_p.add_run(f"  {p.get('code', '')}  |  {p.get('name', '')}")
    _apply_font(tag_run, size=10, color=COLOR_DARK_GRAY)

    # ══════════════════════════════════════════════════════════════════
    # 4. DOCUMENT TITLE (Large, Left-aligned — corporate style)
    # ══════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after  = Pt(4)
    title_run = title_p.add_run(doc_title)
    _apply_font(title_run, size=28, bold=True, color=COLOR_PRIMARY)

    # Underscore accent line under title
    _add_bottom_border(title_p)

    # Subtitle (Thai)
    if doc_subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_p.paragraph_format.space_before = Pt(4)
        sub_p.paragraph_format.space_after  = Pt(20)
        sub_run = sub_p.add_run(doc_subtitle)
        _apply_font(sub_run, size=13, italic=True, color=COLOR_DARK_GRAY)
    else:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(20)

    # ══════════════════════════════════════════════════════════════════
    # 5. INFO TABLE — Two-column, styled
    # ══════════════════════════════════════════════════════════════════
    info_data = [
        ("โครงการ / Project",              p.get("name", "")),
        ("รหัสเอกสาร / Document ID",       doc_id),
        ("เวอร์ชัน / Version",             p.get("version", "1.0")),
        ("วันที่จัดทำ / Document Date",    p.get("document_date", datetime.now().strftime("%Y-%m-%d"))),
        ("หน่วยงาน / Department",          p.get("department", "")),
        ("ระดับความลับ / Classification",  p.get("classification", "Internal")),
    ]

    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        # Uniform gray label column (consistent, clean)
        _set_cell_text(row.cells[0], label, bold=True,
                       bg_color=COLOR_SECONDARY, font_color=COLOR_BLACK, size=SZ_TABLE)
        # Alternating very light gray on value column for readability
        val_bg = COLOR_LIGHT_GRAY if i % 2 == 1 else None
        _set_cell_text(row.cells[1], value, size=SZ_TABLE, bg_color=val_bg)
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)

    # ══════════════════════════════════════════════════════════════════
    # 6. BOTTOM BRANDED BAR — Navy with ISO note embedded as white text
    # ══════════════════════════════════════════════════════════════════
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(16)
    sp2.paragraph_format.space_after  = Pt(0)

    # Single-cell full-width table with navy background
    bot_tbl = doc.add_table(rows=1, cols=1)
    bot_tbl.style = "Table Grid"
    _remove_table_borders(bot_tbl)
    bot_cell = bot_tbl.rows[0].cells[0]
    bot_cell.width = Cm(16)
    _set_cell_bg(bot_cell, COLOR_PRIMARY)

    bp = bot_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(7)
    bp.paragraph_format.space_after  = Pt(7)
    iso_run = bp.add_run(
        "จัดทำตามมาตรฐาน ISO/IEC 29110  |  Software Engineering — Lifecycle Profiles for VSE"
    )
    _apply_font(iso_run, size=9, italic=True, color=COLOR_WHITE)

    # Page break after cover
    doc.add_page_break()

    # Add header/footer to all pages (called once after cover page)
    add_header_footer(doc, config, doc_id, doc_title)


# ─── Document Control Table ───────────────────────────────────────────────────
def add_document_control(doc: Document, config: dict, doc_id: str,
                          reviewers: list = None, approvers: list = None):
    """Add Document Control section with Prepared/Reviewed/Approved by table."""
    add_section_heading(doc, "Document Control / การควบคุมเอกสาร", level=1)

    t = config["team"]
    pm   = t.get("project_manager", {})
    ba   = t.get("business_analyst", {})
    qa   = t.get("qa_engineer", {})

    if reviewers is None:
        reviewers = [{"name": qa.get("name", ""), "title": qa.get("title", "QA Engineer")}]
    if approvers is None:
        approvers = [{"name": pm.get("name", ""), "title": pm.get("title", "Project Manager")}]

    rows = [
        ["เตรียมโดย / Prepared by",
         ba.get("name", pm.get("name", "—")),
         ba.get("title", "Business Analyst"),
         config["project"].get("document_date", "")],
        *[["ตรวจสอบโดย / Reviewed by", r.get("name",""), r.get("title",""), ""]
          for r in reviewers],
        *[["อนุมัติโดย / Approved by",  a.get("name",""), a.get("title",""), ""]
          for a in approvers],
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["บทบาท / Role", "ชื่อ / Name", "ตำแหน่ง / Title", "วันที่ / Date"]):
        # ISO: table header — light gray bg, black bold text, centered
        _set_cell_text(cell, txt, bold=True, bg_color=COLOR_SECONDARY,
                       font_color=COLOR_BLACK, center=True, size=SZ_TABLE)
    for row_data in rows:
        row = tbl.add_row()
        for cell, txt in zip(row.cells, row_data):
            _set_cell_text(cell, txt, size=SZ_TABLE)

    doc.add_paragraph()


# ─── Version History ──────────────────────────────────────────────────────────
def add_version_history(doc: Document, config: dict):
    """Add Version History table."""
    add_section_heading(doc, "Version History / ประวัติการแก้ไข", level=1)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    headers = ["Version", "วันที่ / Date", "แก้ไขโดย / Author", "ผู้อนุมัติ / Approved", "รายละเอียด / Description"]
    for cell, txt in zip(tbl.rows[0].cells, headers):
        _set_cell_text(cell, txt, bold=True, bg_color=COLOR_SECONDARY,
                       font_color=COLOR_BLACK, center=True, size=SZ_TABLE)

    versions = config.get("versions", [])
    if versions:
        for v in versions:
            row = tbl.add_row()
            vals = [v.get("version","1.0"), v.get("release_date",""),
                    v.get("deployed_by",""), "", v.get("description","Initial release")]
            for cell, txt in zip(row.cells, vals):
                _set_cell_text(cell, txt, size=SZ_TABLE)
    else:
        row = tbl.add_row()
        vals = ["1.0", config["project"].get("document_date", datetime.now().strftime("%Y-%m-%d")),
                config["team"].get("business_analyst",{}).get("name",""),
                config["team"].get("project_manager",{}).get("name",""),
                "Initial version"]
        for cell, txt in zip(row.cells, vals):
            _set_cell_text(cell, txt, size=SZ_TABLE)

    doc.add_paragraph()


# ─── Section Headings ─────────────────────────────────────────────────────────
def add_section_heading(doc: Document, text: str, level: int = 1):
    """
    Add ISO-standard section heading.
    H1: 16pt Bold #003366 + bottom border
    H2: 14pt Bold #003366
    H3: 12pt Bold #003366
    """
    size_map = {1: SZ_H1, 2: SZ_H2, 3: SZ_H3}
    space_before_map = {1: Pt(20), 2: Pt(14), 3: Pt(10)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = space_before_map.get(level, Pt(10))
    p.paragraph_format.space_after  = Pt(6) if level == 1 else Pt(4)

    # H1: add small left navy bar accent via shading a thin table
    if level == 1:
        run = p.add_run(text)
        _apply_font(run, size=size_map[1], bold=True, color=COLOR_PRIMARY)
        _add_bottom_border(p)
    elif level == 2:
        # Prefix with small gold square
        prefix = p.add_run("▌ ")
        _apply_font(prefix, size=size_map[2], bold=True, color=COLOR_ACCENT)
        run = p.add_run(text)
        _apply_font(run, size=size_map[2], bold=True, color=COLOR_PRIMARY)
    else:
        run = p.add_run(text)
        _apply_font(run, size=size_map.get(level, SZ_H3), bold=True, color=COLOR_PRIMARY)

    return p


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                  color: RGBColor = None, size: int = None, indent: float = 0):
    """Add a styled body paragraph (Cambria 11pt by default)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _apply_font(run, size=size or SZ_BODY, bold=bold, italic=italic,
                color=color or COLOR_BLACK)
    return p


BULLET_SYMBOLS = {1: "•", 2: "○", 3: "▪"}   # ISO §9.1 bullet hierarchy

def add_bullet(doc: Document, text: str, level: int = 1):
    """
    Add an ISO-standard bullet point.
    Level 1: •  Level 2: ○  Level 3: ▪
    """
    symbol = BULLET_SYMBOLS.get(level, "•")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(level * 0.75)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{symbol}  {text}")
    _apply_font(run, size=SZ_BODY, color=COLOR_BLACK)
    return p


def add_code_block(doc: Document, code_text: str, caption: str = ""):
    """
    Add a formatted code block per ISO §9.3:
    Courier New 10pt, light gray background (#F2F2F2), thin border.
    Optionally add an italic caption below.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    run = p.add_run(code_text)
    run.font.name  = FONT_CODE
    run.font.size  = Pt(SZ_CODE)
    run.font.color.rgb = COLOR_BLACK

    # Gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)

    if caption:
        add_caption(doc, caption, is_figure=False)
    return p


def add_caption(doc: Document, caption_text: str, is_figure: bool = False):
    """
    Add an ISO-standard table/figure caption.
    Format: 'ตารางที่ N — Description'  or  'รูปที่ N — Description'
    Cambria 10pt Italic per ISO §6.2 / §7.2
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    prefix = "รูปที่" if is_figure else "ตารางที่"
    run = p.add_run(f"{prefix} — {caption_text}")
    _apply_font(run, size=SZ_CAPTION, italic=True, color=COLOR_DARK_GRAY)
    return p


# ─── Tables ───────────────────────────────────────────────────────────────────
def add_table(doc: Document, headers: list, rows: list,
              col_widths: list = None, header_bg: RGBColor = None,
              caption: str = "") -> None:
    """
    Add an ISO-standard formatted table.
    ISO §6.1: header bg #D9D9D9, bold centered; borders 0.5-1pt.
    Args:
        headers   : list of header strings
        rows      : list of lists (row data)
        col_widths: list of widths in cm (optional)
        header_bg : override header bg (default COLOR_SECONDARY = #D9D9D9)
        caption   : optional caption string shown above the table
    """
    # ISO §6.2: caption above table
    if caption:
        add_caption(doc, caption, is_figure=False)

    if header_bg is None:
        header_bg = COLOR_SECONDARY   # #D9D9D9 per ISO standard

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — ISO: bold, centered, light gray bg, black text
    hdr_row = tbl.rows[0]
    for i, (cell, txt) in enumerate(zip(hdr_row.cells, headers)):
        _set_cell_text(cell, txt, bold=True, bg_color=header_bg,
                       font_color=COLOR_BLACK, center=True, size=SZ_TABLE)
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])

    # Data rows — alternating very light gray for readability
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = COLOR_LIGHT_GRAY if ri % 2 == 1 else None
        for i, (cell, txt) in enumerate(zip(row.cells, row_data)):
            _set_cell_text(cell, str(txt) if txt is not None else "",
                           bg_color=bg, size=SZ_TABLE)
            if col_widths and i < len(col_widths):
                cell.width = Cm(col_widths[i])

    doc.add_paragraph()
    return tbl


def add_signature_table(doc: Document, signatories: list):
    """
    Add a signature table at the bottom of a document.
    signatories: list of dicts with keys: name, title, role
    Example: [{"name":"John", "title":"PM", "role":"Approver"}]
    """
    add_section_heading(doc, "ลายมือชื่อ / Signatures", level=1)

    tbl = doc.add_table(rows=3, cols=len(signatories))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Role — ISO gray header
    for i, sig in enumerate(signatories):
        _set_cell_text(tbl.rows[0].cells[i], sig.get("role",""),
                       bold=True, bg_color=COLOR_SECONDARY,
                       font_color=COLOR_BLACK, center=True, size=SZ_TABLE)
    # Row 1: Signature line
    for i in range(len(signatories)):
        cell = tbl.rows[1].cells[i]
        _set_cell_text(cell, "\n\nลงชื่อ ______________________________\n",
                       center=True, size=SZ_BODY)
        cell.height = Cm(3)
    # Row 2: Name + Date
    for i, sig in enumerate(signatories):
        _set_cell_text(tbl.rows[2].cells[i],
                       f"({sig.get('name','')}\n{sig.get('title','')}\nวันที่: _______________",
                       center=True, size=SZ_BODY)
    doc.add_paragraph()


# ─── Private Helpers ──────────────────────────────────────────────────────────
def _apply_font(run, size: int = None, bold: bool = False, italic: bool = False,
               color: RGBColor = None):
    """
    Apply ISO-standard font to a run.
    Sets ascii/hAnsi to Cambria, cs (complex script) to TH Sarabun New for Thai.
    """
    run.font.name  = FONT_BODY
    run.font.size  = Pt(size or SZ_BODY)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color or COLOR_BLACK
    # Set complex-script font via XML so Thai renders with TH Sarabun New
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    FONT_BODY)
    rFonts.set(qn("w:hAnsi"),    FONT_BODY)
    rFonts.set(qn("w:cs"),       FONT_THAI)
    rFonts.set(qn("w:eastAsia"), FONT_THAI)


def _set_cell_text(cell, text: str, bold: bool = False, bg_color: RGBColor = None,
                   font_color: RGBColor = None, center: bool = False, size: int = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _apply_font(run, size=size or SZ_TABLE, bold=bold, color=font_color or COLOR_BLACK)
    if bg_color:
        _set_cell_bg(cell, bg_color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _rgb_hex(color: RGBColor) -> str:
    """Convert RGBColor to hex string. Works with both old (attributes) and new (tuple) API."""
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def _set_cell_bg(cell, color: RGBColor):
    """Set cell background color via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = _rgb_hex(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_color_bar(doc: Document, color: RGBColor, height_cm: float = 0.5):
    """Add a colored horizontal bar (paragraph with background shading)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(height_cm * 28.35)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = _rgb_hex(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _add_bottom_border(paragraph):
    """Add a bottom border to a paragraph (used for section headings)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _rgb_hex(COLOR_PRIMARY))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_para_bottom_border(paragraph, color_hex: str = "1F497D"):
    """Add a bottom border to a paragraph (used in header/footer)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_para_top_border(paragraph, color_hex: str = "003366"):
    """Add a top border to a paragraph (used above footer)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def _remove_table_borders(table):
    """Remove all borders from a table (for layout tables in header/cover)."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _add_page_number(paragraph):
    """Insert 'Page X of Y' field into a paragraph (Cambria 10pt)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("หน้าที่ ")
    _apply_font(run, size=SZ_HEADER, color=COLOR_DARK_GRAY)

    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = paragraph.add_run()
    _apply_font(run2, size=SZ_HEADER, color=COLOR_DARK_GRAY)
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    run3 = paragraph.add_run(" / ")
    _apply_font(run3, size=SZ_HEADER, color=COLOR_DARK_GRAY)

    # NUMPAGES field
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run4 = paragraph.add_run()
    _apply_font(run4, size=SZ_HEADER, color=COLOR_DARK_GRAY)
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)


def get_doc_id(config: dict, folder_num: str, doc_code: str) -> str:
    """Build standard Document ID: [ProjectCode]-[FolderNum]-[DocCode]-v[Version]"""
    code    = config["project"].get("code", "PROJ")
    version = config["project"].get("version", "1.0")
    return f"{code}-{folder_num}-{doc_code}-v{version}"
