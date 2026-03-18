"""
Template: Statement of Work (SOW)
Folder 01 - Project Management
Document ID pattern: [CODE]-01-SOW-v[VER]
ISO/IEC 29110 Reference: PM Process - Project Initiation / Contract Management

Sections:
  1. ข้อมูลสัญญา / Contract Overview
  2. ขอบเขตงาน / Scope of Work
  3. งานที่ส่งมอบ / Deliverables
  4. กำหนดการ / Project Timeline  (4.1 Milestone table + 4.2 Gantt Chart)
  5. ความรับผิดชอบ / Responsibilities
  6. เงื่อนไขการชำระเงิน / Payment Terms
  7. ข้อกำหนดเพิ่มเติม / Terms & Conditions
  8. ลายมือชื่อ / Sign-off
"""
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from datetime import datetime
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from utils.doc_builder import (
    new_document, add_cover_page, add_document_control,
    add_version_history, add_section_heading, add_paragraph,
    add_table, add_signature_table, get_doc_id,
)


# ─── Gantt Chart Helpers ──────────────────────────────────────────────────────

_MONTH_TH = ["ม.ค.", "ก.พ.", "มี.ค.", "เม.ย.", "พ.ค.", "มิ.ย.",
             "ก.ค.", "ส.ค.", "ก.ย.", "ต.ค.", "พ.ย.", "ธ.ค."]

_STATUS_COLOR = {
    "Completed":   "378610",   # Green
    "In Progress": "0070C0",   # Blue
    "Mitigated":   "378610",
    "Cancelled":   "808080",
}
_DEFAULT_BAR = "003366"    # Navy — Planned
_HEADER_BG   = "003366"
_HEADER_TEXT = "FFFFFF"
_LABEL_BG    = "EEF2F7"
_LABEL_TEXT  = "003366"
_INACTIVE_BG = "F2F2F2"
_TODAY_BG    = "FFC000"    # Gold
_LEGEND_ITEMS = [
    ("378610", "เสร็จแล้ว (Completed)"),
    ("0070C0", "กำลังดำเนินการ (In Progress)"),
    ("003366", "แผน (Planned)"),
    ("FFC000", "เดือนปัจจุบัน (Today)"),
]


def _gantt_months(start_str: str, end_str: str) -> list:
    """Return list of (year, month) tuples from start_str to end_str inclusive."""
    try:
        s = datetime.strptime(start_str[:7], "%Y-%m")
        e = datetime.strptime(end_str[:7],   "%Y-%m")
        result, y, m = [], s.year, s.month
        while (y, m) <= (e.year, e.month):
            result.append((y, m))
            m += 1
            if m > 12:
                m, y = 1, y + 1
        return result
    except Exception:
        return []


def _gantt_set_bg(cell, hex_color: str):
    """Set cell fill colour from a 6-char hex string (no #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _gantt_text(cell, text: str, size_pt: int = 8, bold: bool = False,
                fg_hex: str = "000000", center: bool = True):
    """Write text into a Gantt cell with ISO fonts."""
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name  = "Cambria"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.color.rgb = RGBColor(
        int(fg_hex[0:2], 16), int(fg_hex[2:4], 16), int(fg_hex[4:6], 16)
    )
    rPr    = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    "Cambria")
    rFonts.set(qn("w:hAnsi"),    "Cambria")
    rFonts.set(qn("w:cs"),       "TH Sarabun New")
    rFonts.set(qn("w:eastAsia"), "TH Sarabun New")


def _gantt_rotate(cell):
    """Rotate cell text 90° (btLr) for month header columns."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    td   = OxmlElement("w:textDirection")
    td.set(qn("w:val"), "btLr")
    tcPr.append(td)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


def _gantt_row_height(row, height_cm: float):
    """Lock row to exact height in cm."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(height_cm * 567)))  # 567 twips ≈ 1 cm
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _add_gantt_chart(doc, config: dict):
    """
    Render a table-based Gantt chart into the document.

    Layout:
      Row 0   : Month header labels (rotated 90°, navy background)
      Row 1…N : One row per milestone — coloured bar spans active months

    Colour coding:
      ■ Completed   → green   #378610
      ■ In Progress → blue    #0070C0
      ■ Planned     → navy    #003366
      ■ Current mth → gold    #FFC000
      □ Inactive    → lt-gray #F2F2F2
    """
    milestones = config.get("milestones", [])
    p          = config["project"]
    start_str  = p.get("start_date", "")
    end_str    = p.get("end_date",   "")
    months     = _gantt_months(start_str, end_str)

    if not milestones or not months:
        add_paragraph(doc, "— ไม่มีข้อมูล Milestone สำหรับสร้าง Gantt Chart —")
        return

    today_ym = (datetime.today().year, datetime.today().month)
    n_months = len(months)

    LABEL_W = 4.5
    month_w = min(1.2, (16.0 - LABEL_W) / n_months)

    # Pre-compute bar spans: milestone i spans from end of (i-1) to its own target_date
    rows_data = []
    for i, ms in enumerate(milestones):
        prev_end  = start_str if i == 0 else milestones[i - 1].get("target_date", start_str)
        this_end  = ms.get("target_date", end_str)
        status    = ms.get("status", "Planned")
        bar_color = _STATUS_COLOR.get(status, _DEFAULT_BAR)
        rows_data.append({
            "label":     f"{ms.get('id', '')}  {ms.get('name', '')}",
            "start_str": prev_end,
            "end_str":   this_end,
            "bar_color": bar_color,
        })

    # ── Build table: (1 + N) rows × (1 + M) columns ──────────────────────────
    tbl = doc.add_table(rows=1 + len(rows_data), cols=1 + n_months)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl._tblPr.append(layout)

    # ── Header row ────────────────────────────────────────────────────────────
    hdr = tbl.rows[0]
    _gantt_row_height(hdr, 1.8)

    hdr.cells[0].width = Cm(LABEL_W)
    _gantt_set_bg(hdr.cells[0], _HEADER_BG)
    _gantt_text(hdr.cells[0], "Milestone / กิจกรรม",
                size_pt=9, bold=True, fg_hex=_HEADER_TEXT)

    for j, (yr, mo) in enumerate(months):
        cell       = hdr.cells[j + 1]
        cell.width = Cm(month_w)
        is_now     = (yr, mo) == today_ym
        _gantt_set_bg(cell, _TODAY_BG if is_now else _HEADER_BG)
        _gantt_rotate(cell)
        _gantt_text(cell,
                    f"{_MONTH_TH[mo - 1]}'{str(yr)[2:]}",
                    size_pt=8, bold=True,
                    fg_hex="003366" if is_now else _HEADER_TEXT)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for i, rd in enumerate(rows_data):
        row = tbl.rows[i + 1]
        _gantt_row_height(row, 0.65)

        try:
            ms_start = datetime.strptime(rd["start_str"][:7], "%Y-%m")
            ms_end   = datetime.strptime(rd["end_str"][:7],   "%Y-%m")
        except Exception:
            ms_start = ms_end = None

        lbl       = row.cells[0]
        lbl.width = Cm(LABEL_W)
        _gantt_set_bg(lbl, _LABEL_BG)
        _gantt_text(lbl, rd["label"], size_pt=8, bold=True,
                    fg_hex=_LABEL_TEXT, center=False)
        lbl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for j, (yr, mo) in enumerate(months):
            cell       = row.cells[j + 1]
            cell.width = Cm(month_w)
            is_now     = (yr, mo) == today_ym
            cur_dt     = datetime(yr, mo, 1)

            in_bar = ms_start and ms_end and ms_start <= cur_dt <= ms_end
            if in_bar:
                fill = _TODAY_BG if is_now else rd["bar_color"]
            else:
                fill = _TODAY_BG if is_now else _INACTIVE_BG
            _gantt_set_bg(cell, fill)

    doc.add_paragraph()

    # ── Legend ────────────────────────────────────────────────────────────────
    leg = doc.add_table(rows=1, cols=len(_LEGEND_ITEMS) * 2)
    leg.style     = "Table Grid"
    leg.alignment = WD_TABLE_ALIGNMENT.LEFT
    _gantt_row_height(leg.rows[0], 0.55)

    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    leg._tblPr.append(tblBdr)

    for k, (hex_c, label_th) in enumerate(_LEGEND_ITEMS):
        swatch     = leg.rows[0].cells[k * 2]
        label_cell = leg.rows[0].cells[k * 2 + 1]
        swatch.width     = Cm(0.45)
        label_cell.width = Cm(3.3)
        _gantt_set_bg(swatch, hex_c)
        _gantt_text(label_cell, f"  {label_th}", size_pt=8, fg_hex="404040", center=False)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


def generate(config: dict, output_path: str) -> str:
    """Generate Statement of Work .docx and return saved file path."""
    doc_id = get_doc_id(config, "01", "SOW")
    doc    = new_document()
    p      = config["project"]
    t      = config["team"]
    sow    = config.get("sow", {})

    # Fallback defaults when no sow section provided
    client = sow.get("client", {})
    vendor = sow.get("vendor", {})

    # Cover
    add_cover_page(doc, config, doc_id, "Statement of Work", "ขอบเขตงานและสัญญาว่าจ้าง")

    # Document Control
    add_document_control(doc, config, doc_id)

    # Version History
    add_version_history(doc, config)

    doc.add_page_break()

    # ── 1. Contract Overview ──────────────────────────────────────────────────
    add_section_heading(doc, "1. ข้อมูลสัญญา / Contract Overview", level=1)
    add_table(doc,
        headers=["หัวข้อ / Field", "รายละเอียด / Details"],
        rows=[
            ["เลขที่สัญญา / Contract No.",       sow.get("contract_number", "—")],
            ["วันที่มีผล / Effective Date",        sow.get("effective_date", p.get("start_date", "—"))],
            ["วันที่สิ้นสุด / Expiration Date",   sow.get("expiration_date", p.get("end_date", "—"))],
            ["ชื่อโครงการ / Project Name",         p.get("name", "—")],
            ["รหัสโครงการ / Project Code",         p.get("code", "—")],
            ["หน่วยงาน / Department",              p.get("department", "—")],
        ],
        col_widths=[6, 10]
    )

    add_section_heading(doc, "1.1 ฝ่ายว่าจ้าง / Client", level=2)
    add_table(doc,
        headers=["หัวข้อ", "รายละเอียด"],
        rows=[
            ["ชื่อองค์กร / Organization",  client.get("name", p.get("organization", "—"))],
            ["ผู้แทน / Representative",    client.get("representative", "—")],
            ["ตำแหน่ง / Title",            client.get("title", "—")],
            ["อีเมล / Email",              client.get("email", "—")],
            ["โทรศัพท์ / Phone",           client.get("phone", "—")],
        ],
        col_widths=[6, 10]
    )

    add_section_heading(doc, "1.2 ผู้รับจ้าง / Vendor", level=2)
    add_table(doc,
        headers=["หัวข้อ", "รายละเอียด"],
        rows=[
            ["ชื่อองค์กร / Organization",  vendor.get("name", p.get("organization", "—"))],
            ["ผู้แทน / Representative",    vendor.get("representative", t.get("project_manager", {}).get("name", "—"))],
            ["ตำแหน่ง / Title",            vendor.get("title", "Project Manager")],
            ["อีเมล / Email",              vendor.get("email", t.get("project_manager", {}).get("email", "—"))],
            ["โทรศัพท์ / Phone",           vendor.get("phone", "—")],
        ],
        col_widths=[6, 10]
    )

    # ── 2. Scope of Work ─────────────────────────────────────────────────────
    add_section_heading(doc, "2. ขอบเขตงาน / Scope of Work", level=1)
    add_paragraph(doc, p.get("description", "—"))

    add_section_heading(doc, "2.1 ขอบเขตโครงการ / Project Scope", level=2)
    scope_text = p.get("scope", "—")
    for line in scope_text.split("\n"):
        line = line.strip()
        if line:
            add_paragraph(doc, f"• {line}" if not line.startswith("•") else line)

    add_section_heading(doc, "2.2 วัตถุประสงค์ / Objectives", level=2)
    for obj in p.get("objectives", ["—"]):
        add_paragraph(doc, f"• {obj}")

    add_section_heading(doc, "2.3 รายการฟีเจอร์ / Feature List", level=2)
    add_paragraph(doc, "รายการฟีเจอร์/ฟังก์ชันหลักที่ครอบคลุมในขอบเขตงานนี้:")
    requirements = config.get("requirements", [])
    if requirements:
        # Group by category for readability
        from collections import OrderedDict
        categories = OrderedDict()
        for req in requirements:
            cat = req.get("category", "General")
            categories.setdefault(cat, []).append(req)

        feat_rows = []
        for cat, reqs in categories.items():
            for req in reqs:
                feat_rows.append([
                    req.get("id", ""),
                    cat,
                    req.get("title", ""),
                    req.get("priority", "Medium"),
                    req.get("type", "Functional"),
                ])
        add_table(doc,
            headers=["REQ ID", "หมวด / Category", "ฟีเจอร์ / Feature",
                     "ความสำคัญ / Priority", "ประเภท / Type"],
            rows=feat_rows,
            col_widths=[2, 3.5, 6, 2.5, 2.5]
        )
    else:
        # Fallback to design_components if no requirements
        components = config.get("design_components", [])
        if components:
            comp_rows = [
                [c.get("id", ""), c.get("name", ""), c.get("description", ""), c.get("type", "")]
                for c in components
            ]
            add_table(doc,
                headers=["ID", "Component / Module", "คำอธิบาย / Description", "ประเภท / Type"],
                rows=comp_rows,
                col_widths=[2, 4.5, 7.5, 2.5]
            )
        else:
            add_paragraph(doc, "— ยังไม่มีรายการฟีเจอร์ โปรดระบุใน requirements หรือ design_components ในไฟล์ config —")

    # ── 3. Deliverables ──────────────────────────────────────────────────────
    add_section_heading(doc, "3. งานที่ส่งมอบ / Deliverables", level=1)
    deliverables = sow.get("deliverables", [])
    if deliverables:
        del_rows = [
            [
                d.get("id", ""),
                d.get("description", ""),
                d.get("due_date", ""),
                d.get("acceptance_criteria", ""),
                d.get("payment_milestone", ""),
            ]
            for d in deliverables
        ]
    else:
        del_rows = [["DEL-001", "ระบุ deliverable ที่นี่", p.get("go_live_date", "—"), "ระบุเงื่อนไขการตรวจรับ", "MS-05"]]

    add_table(doc,
        headers=["ID", "รายละเอียด / Description", "กำหนดส่ง / Due Date",
                 "เกณฑ์ตรวจรับ / Acceptance Criteria", "Milestone"],
        rows=del_rows,
        col_widths=[2, 5.5, 2.5, 5.5, 2]
    )

    # ── 4. Timeline ──────────────────────────────────────────────────────────
    add_section_heading(doc, "4. กำหนดการโครงการ / Project Timeline", level=1)

    add_section_heading(doc, "4.1 ตาราง Milestone", level=2)
    milestones = config.get("milestones", [])
    ms_rows = (
        [[ms.get("id", ""), ms.get("name", ""), ms.get("target_date", ""),
          ms.get("owner", ""), ms.get("status", "Planned")]
         for ms in milestones]
        if milestones else [["MS-01", "—", "—", "—", "Planned"]]
    )
    add_table(doc,
        headers=["ID", "Milestone", "Target Date", "Owner", "Status"],
        rows=ms_rows,
        col_widths=[2.5, 7, 3, 4, 3]
    )

    add_section_heading(doc, "4.2 แผนภูมิ Gantt / Gantt Chart", level=2)
    _add_gantt_chart(doc, config)

    # ── 5. Responsibilities ──────────────────────────────────────────────────
    add_section_heading(doc, "5. ความรับผิดชอบ / Responsibilities", level=1)

    client_resp = sow.get("client_responsibilities", [
        "ให้ข้อมูล Business Requirements และ Domain Knowledge",
        "มอบหมาย Stakeholder หลักสำหรับ Requirement Sessions และ UAT",
        "จัดเตรียม Test Data และสิทธิ์การเข้าถึง Test Environment",
        "ตรวจรับและลงนามเอกสาร Deliverable ภายใน 5 วันทำการ",
    ])
    vendor_resp = sow.get("vendor_responsibilities", [
        "จัดส่งทีมงานที่มีคุณสมบัติเหมาะสมตลอดระยะเวลาโครงการ",
        "ส่งมอบ Deliverables ทั้งหมดตามกำหนดการที่ตกลงไว้",
        "ปฏิบัติตาม Coding Standards และข้อกำหนดเอกสาร ISO/IEC 29110",
        "จัดทำ Status Report รายสัปดาห์ให้กับ Client",
    ])

    max_rows = max(len(client_resp), len(vendor_resp))
    client_resp_padded = client_resp + [""] * (max_rows - len(client_resp))
    vendor_resp_padded = vendor_resp + [""] * (max_rows - len(vendor_resp))
    resp_rows = [
        [f"• {c}" if c else "", f"• {v}" if v else ""]
        for c, v in zip(client_resp_padded, vendor_resp_padded)
    ]
    add_table(doc,
        headers=["ฝ่ายว่าจ้าง (Client)", "ผู้รับจ้าง (Vendor)"],
        rows=resp_rows,
        col_widths=[8, 8]
    )

    # ── 6. Payment Terms ─────────────────────────────────────────────────────
    add_section_heading(doc, "6. เงื่อนไขการชำระเงิน / Payment Terms", level=1)
    payment_terms = sow.get("payment_terms",
        "ชำระเงินภายใน 30 วันหลังได้รับใบแจ้งหนี้ โดยออกใบแจ้งหนี้เมื่อตรวจรับ Milestone Deliverable แต่ละรายการเป็นที่เรียบร้อยแล้ว")
    add_paragraph(doc, payment_terms)

    payment_schedule = sow.get("payment_schedule", [])
    if payment_schedule:
        pay_rows = [
            [
                ps.get("milestone", ""),
                f"{ps.get('percentage', '')}%",
                ps.get("amount", "—"),
                ps.get("due_date", ""),
            ]
            for ps in payment_schedule
        ]
        add_table(doc,
            headers=["Milestone", "% ของมูลค่าสัญญา", "จำนวนเงิน / Amount", "กำหนดชำระ / Due Date"],
            rows=pay_rows,
            col_widths=[7, 3, 3.5, 4]
        )
    else:
        add_paragraph(doc, "— ระบุตารางชำระเงินที่นี่ —")

    # ── 7. Terms & Conditions ────────────────────────────────────────────────
    add_section_heading(doc, "7. ข้อกำหนดเพิ่มเติม / Terms & Conditions", level=1)
    terms = sow.get("terms_conditions", [
        "การรักษาความลับ (Confidentiality): คู่สัญญาทั้งสองฝ่ายตกลงเก็บรักษาข้อมูลทั้งหมดเป็นความลับ",
        "ทรัพย์สินทางปัญญา (IP): Deliverables ทั้งหมดที่ผลิตภายใต้ SOW นี้ตกเป็นกรรมสิทธิ์ของ Client เมื่อชำระเงินครบถ้วน",
        "การจัดการการเปลี่ยนแปลง (Change Management): การเปลี่ยนแปลง Scope ต้องจัดทำ Change Request Form (CRF) และลงนามโดยทั้งสองฝ่าย",
        "การบอกเลิกสัญญา (Termination): คู่สัญญาฝ่ายใดฝ่ายหนึ่งสามารถบอกเลิกสัญญาได้โดยแจ้งล่วงหน้าเป็นลายลักษณ์อักษร 30 วัน",
        "กฎหมายที่ใช้บังคับ (Governing Law): สัญญานี้อยู่ภายใต้กฎหมายไทย",
    ])
    for i, term in enumerate(terms, 1):
        add_paragraph(doc, f"{i}. {term}")

    # ── 8. Sign-off ──────────────────────────────────────────────────────────
    doc.add_page_break()
    add_section_heading(doc, "8. ลายมือชื่อ / Sign-off", level=1)
    add_paragraph(doc,
        "คู่สัญญาทั้งสองฝ่ายได้อ่านและเข้าใจข้อกำหนดทั้งหมดในเอกสารนี้ และตกลงยินยอมผูกพันตามเงื่อนไขที่ระบุไว้\n"
        "Both parties have read, understood, and agreed to be bound by the terms of this Statement of Work.")

    pm = t.get("project_manager", {})
    add_signature_table(doc, [
        {
            "role":  "ผู้แทนฝ่ายว่าจ้าง / Client Representative",
            "name":  client.get("representative", "—"),
            "title": client.get("title", "—"),
        },
        {
            "role":  "ผู้แทนผู้รับจ้าง / Vendor Representative",
            "name":  vendor.get("representative", pm.get("name", "—")),
            "title": vendor.get("title", pm.get("title", "Project Manager")),
        },
        {
            "role":  "Project Manager",
            "name":  pm.get("name", "—"),
            "title": pm.get("title", "Project Manager"),
        },
    ])

    os.makedirs(output_path, exist_ok=True)
    file_path = os.path.join(output_path, f"{doc_id}_Statement_of_Work.docx")
    doc.save(file_path)
    return file_path
