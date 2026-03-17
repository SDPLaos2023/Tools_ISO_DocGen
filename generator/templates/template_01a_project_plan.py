"""
Template: Project Plan
Folder 01 - Project Management
Document ID pattern: [CODE]-01-PP-v[VER]
ISO/IEC 29110 Reference: PM Process - Project Plan
"""
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.doc_builder import (
    new_document, add_cover_page, add_document_control,
    add_version_history, add_section_heading, add_paragraph,
    add_table, add_signature_table, get_doc_id, COLOR_PRIMARY, COLOR_SECONDARY
)


def generate(config: dict, output_path: str) -> str:
    """Generate Project Plan .docx and return saved file path."""
    doc_id    = get_doc_id(config, "01", "PP")
    doc_title = "Project Plan\nแผนโครงการ"
    doc       = new_document()
    p         = config["project"]
    t         = config["team"]

    # Cover
    add_cover_page(doc, config, doc_id, "Project Plan", "แผนโครงการ")

    # Document Control
    add_document_control(doc, config, doc_id)

    # Version History
    add_version_history(doc, config)

    doc.add_page_break()

    # ── 1. Project Overview ──────────────────────────────────────────────────
    add_section_heading(doc, "1. ภาพรวมโครงการ / Project Overview", level=1)
    add_table(doc,
        headers=["หัวข้อ / Field", "รายละเอียด / Details"],
        rows=[
            ["ชื่อโครงการ / Project Name",        p.get("name","")],
            ["รหัสโครงการ / Project Code",        p.get("code","")],
            ["คำอธิบาย / Description",            p.get("description","")],
            ["ขอบเขต / Scope",                    p.get("scope","")],
            ["หน่วยงาน / Organization",           p.get("organization","")],
            ["ฝ่าย / Department",                 p.get("department","")],
            ["วันเริ่ม / Start Date",             p.get("start_date","")],
            ["วันสิ้นสุด / End Date",             p.get("end_date","")],
            ["วัน Go-Live",                       p.get("go_live_date","")],
        ],
        col_widths=[6, 10]
    )

    # ── 2. Objectives ────────────────────────────────────────────────────────
    add_section_heading(doc, "2. วัตถุประสงค์ / Objectives", level=1)
    for obj in p.get("objectives", ["—"]):
        add_paragraph(doc, f"• {obj}")

    # ── 3. Team & Roles ──────────────────────────────────────────────────────
    add_section_heading(doc, "3. ทีมงานและบทบาท / Team & Roles", level=1)
    team_rows = []
    for role_key, label_th in [
        ("project_manager","Project Manager"),
        ("lead_developer","Lead Developer"),
        ("business_analyst","Business Analyst"),
        ("system_analyst","System Analyst"),
        ("qa_engineer","QA Engineer"),
        ("dba","DBA"),
    ]:
        member = t.get(role_key, {})
        if member.get("name"):
            team_rows.append([label_th, member.get("name",""), member.get("email","")])
    for m in t.get("members", []):
        team_rows.append([m.get("title","Member"), m.get("name",""), m.get("email","")])

    add_table(doc,
        headers=["บทบาท / Role", "ชื่อ / Name", "อีเมล / Email"],
        rows=team_rows,
        col_widths=[5, 6, 7]
    )

    # ── 4. Stakeholders ──────────────────────────────────────────────────────
    add_section_heading(doc, "4. ผู้มีส่วนได้เสีย / Stakeholders", level=1)
    stake_rows = [[s.get("name",""), s.get("role",""), s.get("organization",""), s.get("responsibility","")]
                  for s in config.get("stakeholders", [])]
    if not stake_rows:
        stake_rows = [["—","—","—","—"]]
    add_table(doc,
        headers=["ชื่อ / Name","บทบาท / Role","หน่วยงาน / Org","ความรับผิดชอบ / Responsibility"],
        rows=stake_rows,
        col_widths=[5,4,5,7]
    )

    # ── 5. Project Schedule / Milestones ─────────────────────────────────────
    add_section_heading(doc, "5. กำหนดการโครงการ / Project Schedule", level=1)
    milestone_rows = [
        [ms.get("id",""), ms.get("name",""), ms.get("target_date",""),
         ms.get("owner",""), ms.get("status","Planned")]
        for ms in config.get("milestones", [])
    ]
    if not milestone_rows:
        milestone_rows = [["MS-01","—","—","—","Planned"]]
    add_table(doc,
        headers=["ID","Milestone","Target Date","Owner","Status"],
        rows=milestone_rows,
        col_widths=[2.5, 7, 3, 4, 3]
    )

    # ── 6. Technology Stack ───────────────────────────────────────────────────
    add_section_heading(doc, "6. เทคโนโลยีที่ใช้ / Technology Stack", level=1)
    ts = config.get("tech_stack", {})
    tech_rows = [
        ["Frontend",      ts.get("frontend","—")],
        ["Backend",       ts.get("backend","—")],
        ["Database",      ts.get("database","—")],
        ["Infrastructure",ts.get("infrastructure","—")],
        ["Source Control",ts.get("source_control","—")],
        ["CI/CD",         ts.get("ci_cd","—")],
    ]
    for other in ts.get("other", []):
        tech_rows.append(["Other", other])
    add_table(doc,
        headers=["ประเภท / Category","รายละเอียด / Detail"],
        rows=tech_rows,
        col_widths=[6, 10]
    )

    # ── 7. Risk Summary ───────────────────────────────────────────────────────
    add_section_heading(doc, "7. ความเสี่ยงหลัก / Key Risks (Summary)", level=1)
    add_paragraph(doc, "รายละเอียดเต็มอยู่ใน Risk Register (เอกสาร 09-RR)")
    risk_rows = [
        [r.get("id",""), r.get("description",""), r.get("risk_level",""), r.get("mitigation","")]
        for r in config.get("risks", [])[:5]
    ]
    if not risk_rows:
        risk_rows = [["RISK-001","—","Medium","—"]]
    add_table(doc,
        headers=["ID","ความเสี่ยง / Risk","ระดับ / Level","แผนลด / Mitigation"],
        rows=risk_rows,
        col_widths=[2.5, 7, 3, 7]
    )

    # ── 8. Assumptions & Constraints ─────────────────────────────────────────
    add_section_heading(doc, "8. ข้อสมมติฐานและข้อจำกัด / Assumptions & Constraints", level=1)
    add_section_heading(doc, "8.1 Assumptions", level=2)
    add_paragraph(doc, "• ทรัพยากรทีมพร้อมใช้งานตามแผน")
    add_paragraph(doc, "• Requirements ได้รับการอนุมัติจาก Stakeholder ก่อนเริ่ม Development")
    add_section_heading(doc, "8.2 Constraints", level=2)
    add_paragraph(doc, "• งบประมาณและระยะเวลาตามที่กำหนดในแผน")
    add_paragraph(doc, "• ต้องเป็นไปตามมาตรฐาน ISO/IEC 29110")

    # ── 9. Sign-off ───────────────────────────────────────────────────────────
    doc.add_page_break()
    add_signature_table(doc, [
        {"role": "เตรียมโดย / Prepared by",     "name": t.get("business_analyst",{}).get("name",""), "title": "Business Analyst"},
        {"role": "ตรวจสอบโดย / Reviewed by",   "name": t.get("qa_engineer",{}).get("name",""),       "title": "QA Engineer"},
        {"role": "อนุมัติโดย / Approved by",    "name": t.get("project_manager",{}).get("name",""),  "title": "Project Manager"},
    ])

    os.makedirs(output_path, exist_ok=True)
    file_path = os.path.join(output_path, f"{doc_id}_Project_Plan.docx")
    doc.save(file_path)
    return file_path
