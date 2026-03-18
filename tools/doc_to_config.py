"""
doc_to_config.py — Reverse-Extract Config from Existing .docx Files
=====================================================================
อ่านไฟล์ .docx ที่อัปเดตมาจากภายนอกจาก projects/[ProjectCode]/
และ rewrite configs/[ProjectCode]/[ProjectCode]_config.json ใหม่

Usage:
    python tools/doc_to_config.py --project Dvets
    python tools/doc_to_config.py --project Dvets --output configs/Dvets/Dvets_config_new.json
    python tools/doc_to_config.py --project Dvets --no-backup
    python tools/doc_to_config.py --project Dvets --dry-run   (show extracted data, no write)

Strategy:
    1. Load existing config as base (preserve emails, output_path, etc.)
    2. Scan project folder for .docx files — map by filename code
    3. Extract structured data from each document
    4. Merge: override extracted sections into existing config
    5. Backup existing config → write new config

Document → Config mapping:
    PP  → project, team, stakeholders, milestones, tech_stack
    MM  → meetings
    SOW → sow
    BRD → requirements
    SDD → design_components
    DBD → database_tables
    TCR → test_cases
    BDL → defects
    RR  → risks
    CRF → change_requests
    VRN → versions (+ deployments)
    TR  → training_sessions
    ISL → incidents
    AR  → audit
    CAPA → capas
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Force UTF-8 output on Windows (avoid cp1252 UnicodeEncodeError)
if sys.platform == "win32":
    import os as _os
    _os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass
    try:
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass

# ─── Workspace root detection ────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
WORKSPACE_ROOT = SCRIPT_DIR.parent  # d:\Work\DocISOGen


# ─── Folder → doc-code map ───────────────────────────────────────────────────
FOLDER_DOC_MAP = {
    "01_Project_Management":    ["PP", "MM", "SOW"],
    "02_Requirements_Analysis": ["BRD", "RTM"],
    "03_Design_Architecture":   ["SDD", "DBD"],
    "04_Development":           ["CS", "CRR"],
    "05_Testing_QA":            ["TP", "TCR", "BDL"],
    "06_Deployment_Training":   ["UM", "TR"],
    "07_Support_Maintenance":   ["ISL"],
    "08_Change_Logs_Versioning": ["CRF", "VRN"],
    "09_Risk_Management":       ["RR"],
    "10_Regulatory_Compliance": ["IC", "AR", "CAPA"],
}


# ─── Low-level helpers ────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    """Get clean text from a table cell."""
    return cell.text.strip()


def _normalize(text: str) -> str:
    """Lowercase, strip, remove special chars for label matching."""
    return re.sub(r"[^a-z0-9ก-๙\s]", "", text.lower()).strip()


def read_kv_table(table) -> dict:
    """
    Parse a 2-column Key-Value table.
    Returns {label_raw: value, ...}
    Handles merged cells gracefully.
    """
    result = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip()
            val = cells[1].text.strip()
            if key:
                result[key] = val
    return result


def read_row_table(table) -> list:
    """
    Parse a multi-column table with a header row.
    Returns list of dicts: [{col_header: cell_value, ...}, ...]
    """
    if not table.rows:
        return []
    headers = [_cell_text(c) for c in table.rows[0].cells]
    rows = []
    for row in table.rows[1:]:
        cells = [_cell_text(c) for c in row.cells]
        if all(c == "" for c in cells):
            continue
        row_dict = {}
        for h, v in zip(headers, cells):
            if h:
                row_dict[h] = v
        rows.append(row_dict)
    return rows


def kv_lookup(kv: dict, *keywords) -> str:
    """
    Find first value in kv dict whose key contains any of the keywords.
    Case-insensitive, strips non-alphanumeric.
    """
    for key, val in kv.items():
        norm = _normalize(key)
        for kw in keywords:
            if kw.lower() in norm:
                return val
    return ""


def find_heading2_blocks(doc) -> list:
    """
    Return list of (heading_text, tables_after_heading) tuples.
    Each entry covers the Heading2 paragraph and any tables before the next heading.
    """
    blocks = []
    current_heading = None
    current_tables = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]  # "p" or "tbl"
        if tag == "p":
            # Check if it's a heading
            style_elem = block.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            if style_elem is not None:
                style_val = style_elem.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", ""
                )
                text = "".join(r.text or "" for r in block.iter(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                )).strip()
                if style_val in ("Heading2", "2") or "Heading 2" in style_val:
                    if current_heading is not None:
                        blocks.append((current_heading, current_tables))
                    current_heading = text
                    current_tables = []
            # Ignore normal paragraphs in block detection
        elif tag == "tbl" and current_heading is not None:
            # Build a temporary Document-like table from the XML element
            from docx.table import Table
            try:
                tbl = Table(block, doc)
                current_tables.append(tbl)
            except Exception:
                pass

    if current_heading is not None:
        blocks.append((current_heading, current_tables))

    return blocks


def find_marker_blocks(doc, pattern: str) -> list:
    """
    Scan document body for section blocks triggered by any paragraph matching `pattern`.
    Also treats Heading 2 styled paragraphs as block starters.
    Returns list of (heading_text, tables_after_heading).
    This replaces find_heading2_blocks() for documents that use [normal] style with
    ▌-prefix or keyword patterns instead of Heading 2.
    """
    regex = re.compile(pattern, re.IGNORECASE)
    blocks = []
    current_heading = None
    current_tables = []
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]
        if tag == "p":
            style_elem = block.find(f".//{{{NS}}}pStyle")
            text = "".join(r.text or "" for r in block.iter(f"{{{NS}}}t")).strip()
            if not text:
                continue
            is_section = False
            if style_elem is not None:
                sv = style_elem.get(f"{{{NS}}}val", "")
                if sv in ("Heading2", "2") or "Heading 2" in sv:
                    is_section = True
            if not is_section and regex.search(text):
                is_section = True
            if is_section:
                if current_heading is not None:
                    blocks.append((current_heading, current_tables))
                current_heading = text
                current_tables = []
        elif tag == "tbl" and current_heading is not None:
            from docx.table import Table as _Table
            try:
                tbl = _Table(block, doc)
                current_tables.append(tbl)
            except Exception:
                pass

    if current_heading is not None:
        blocks.append((current_heading, current_tables))
    return blocks


def extract_id_from_heading(heading_text: str, prefix: str) -> str:
    """
    Extract ID like REQ-001, TC-001, etc. from a heading string.
    Falls back to the full heading text if not found.
    """
    m = re.search(rf"({re.escape(prefix)}-\d+)", heading_text, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return heading_text.strip()


def extract_title_from_heading(heading_text: str, id_str: str) -> str:
    """Remove ID, ▌ prefix, and separators from heading to get the title."""
    title = re.sub(re.escape(id_str), "", heading_text, flags=re.IGNORECASE)
    # Strip leading block chars (▌ and similar), whitespace, dashes, dashes, colons
    title = re.sub(r"^[\s▌▐▍▎▏█\-—–:]+", "", title).strip()
    return title or heading_text.strip()


# ─── Per-document extractors ──────────────────────────────────────────────────

def extract_pp(doc, existing: dict) -> dict:
    """Extract from Project Plan → project, team, stakeholders, milestones, tech_stack."""
    result = {}

    # ---- Find document header KV table (Table 1 style: รหัสเอกสาร, เวอร์ชัน, วันที่จัดทำ, ระดับความลับ)
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if any(k in combined for k in ["รหัสเอกสาร", "document id", "document date", "วันที่จัดทำ"]):
            proj = existing.get("project", {}).copy()
            doc_id_raw = kv_lookup(kv, "document id", "รหัสเอกสาร")
            if doc_id_raw:
                m = re.search(r"-v(\d+\.\d+)", doc_id_raw)
                if m:
                    proj.setdefault("doc_revision", m.group(1))
            doc_date = kv_lookup(kv, "document date", "วันที่จัดทำ")
            if doc_date:
                proj["document_date"] = doc_date
            classif = kv_lookup(kv, "classification", "ระดับความลับ")
            if classif:
                proj["classification"] = classif
            result["project"] = proj
            break

    # ---- Find the project overview KV table (ชื่อโครงการ, รหัสโครงการ, คำอธิบาย, ขอบเขต, หน่วยงาน)
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if any(k in combined for k in ["project name", "ชื่อโครงการ", "project code", "รหัสโครงการ"]):
            proj = result.get("project", existing.get("project", {}).copy())
            proj["name"]          = kv_lookup(kv, "project name", "ชื่อโครงการ") or proj.get("name", "")
            proj["code"]          = kv_lookup(kv, "project code", "รหัสโครงการ") or proj.get("code", "")
            proj["description"]   = kv_lookup(kv, "description", "คำอธิบาย") or proj.get("description", "")
            proj["scope"]         = kv_lookup(kv, "scope", "ขอบเขต") or proj.get("scope", "")
            proj["organization"]  = kv_lookup(kv, "organization", "หน่วยงาน") or proj.get("organization", "")
            proj["department"]    = kv_lookup(kv, "department", "ฝ่าย", "หน่วยงาน") or proj.get("department", "")
            raw_start  = kv_lookup(kv, "start date", "วันเริ่ม")
            raw_end    = kv_lookup(kv, "end date", "วันสิ้นสุด")
            raw_golive = kv_lookup(kv, "go-live", "go live", "golive")
            if raw_start:  proj["start_date"]   = raw_start
            if raw_end:    proj["end_date"]     = raw_end
            if raw_golive: proj["go_live_date"] = raw_golive
            if not proj.get("short_name") and proj.get("code"):
                proj["short_name"] = proj["code"]
            result["project"] = proj
            break

    # ---- Extract objectives from bullet paragraphs after วัตถุประสงค์ section
    in_objectives = False
    objectives = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(kw in text for kw in ["วัตถุประสงค์", "objectives", "2. "]):
            if "วัตถุประสงค์" in text or "objectives" in text.lower():
                in_objectives = True
                continue
        elif in_objectives:
            if re.match(r"^\d+\s*\.", text) and text != "":
                in_objectives = False
                break
            if text.startswith(("•", "▪", "▸", "◆", "-", "–", "*")) and len(text) > 2:
                objectives.append(re.sub(r"^[•▪▸◆\-–*]\s*", "", text).strip())
    if objectives:
        proj = result.get("project", existing.get("project", {}).copy())
        if not proj.get("objectives"):
            proj["objectives"] = objectives
        result["project"] = proj

    # ---- Team table (Role, Name, Email) — prefer table with Email column, not document-control table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if not (any("role" in h or "บทบาท" in h for h in headers_lower) and
                any("name" in h or "ชื่อ" in h for h in headers_lower)):
            continue
        # Skip document-control tables (those with Date column) in favour of Email column tables
        has_email = any("email" in h or "อีเมล" in h for h in headers_lower)
        has_date  = any("date" in h or "วันที่" in h for h in headers_lower)
        if has_date and not has_email:
            continue  # this is the document-control table, skip
        rows = read_row_table(table)
        team = existing.get("team", {}).copy()
        role_key_map = {
            "project manager": "project_manager",
            "pm": "project_manager",
            "lead developer": "lead_developer",
            "head developer": "lead_developer",
            "business analyst": "business_analyst",
            "system analyst": "system_analyst",
            "qa engineer": "qa_engineer",
            "test lead": "qa_engineer",
            "dba": "dba",
            "database administrator": "dba",
        }
        members = []
        for row in rows:
            role_raw  = row.get("บทบาท / Role", row.get("Role", row.get("บทบาท", ""))).strip()
            name_raw  = row.get("ชื่อ / Name",  row.get("Name", row.get("ชื่อ",  ""))).strip()
            email_raw = row.get("อีเมล / Email", row.get("Email", "")).strip()
            if not name_raw:
                continue
            role_norm = _normalize(role_raw)
            matched = False
            for keyword, field in role_key_map.items():
                if keyword in role_norm:
                    existing_member = team.get(field, {})
                    team[field] = {
                        "name":  name_raw,
                        "title": role_raw,
                        "email": email_raw or existing_member.get("email", ""),
                        "phone": existing_member.get("phone", ""),
                    }
                    matched = True
                    break
            if not matched:
                existing_emails = {m.get("name"): m.get("email", "") for m in existing.get("team", {}).get("members", [])}
                members.append({
                    "name": name_raw,
                    "title": role_raw,
                    "email": email_raw or existing_emails.get(name_raw, ""),
                })
        if members:
            team["members"] = members
        result["team"] = team
        break

    # ---- Stakeholders table (Name, Role, Org, Responsibility)
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if len(headers_lower) >= 3 and any("sponsor" in h or "responsibility" in h or "ความรับผิดชอบ" in h for h in headers_lower):
            rows = read_row_table(table)
            stakeholders = []
            ex_stakes = {s.get("name", ""): s for s in existing.get("stakeholders", [])}
            for row in rows:
                name = row.get("ชื่อ / Name", row.get("Name", row.get("ชื่อ", ""))).strip()
                role = row.get("บทบาท / Role", row.get("Role", row.get("บทบาท", ""))).strip()
                org  = row.get("หน่วยงาน / Org", row.get("Org", row.get("หน่วยงาน", row.get("Organization", "")))).strip()
                resp = row.get("ความรับผิดชอบ / Responsibility", row.get("Responsibility", row.get("ความรับผิดชอบ", ""))).strip()
                if not name:
                    continue
                ex = ex_stakes.get(name, {})
                stakeholders.append({
                    "name": name,
                    "role": role,
                    "organization": org or ex.get("organization", ""),
                    "email": ex.get("email", ""),
                    "responsibility": resp,
                })
            if stakeholders:
                result["stakeholders"] = stakeholders
            break

    # ---- Milestones table — require "Milestone" column specifically (not version history)
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if not any("milestone" in h for h in headers_lower):
            continue
        rows = read_row_table(table)
        # Validate: first data row must have an ID starting with "MS-"
        if not rows or not (rows[0].get("ID", rows[0].get("Milestone ID", "")).startswith("MS-") or
                            any(str(v).startswith("MS-") for v in rows[0].values())):
            continue
        milestones = []
        for row in rows:
            ms_id  = row.get("ID", row.get("MS ID", row.get("Milestone ID", ""))).strip()
            name   = row.get("Name", row.get("ชื่อ", row.get("Milestone", ""))).strip()
            date   = row.get("Target Date", row.get("Date", row.get("วันที่", row.get("กำหนด", "")))).strip()
            status = row.get("Status", row.get("สถานะ", "Planned")).strip()
            owner  = row.get("Owner", row.get("เจ้าของ", "")).strip()
            if not ms_id and not name:
                continue
            if not ms_id:
                ms_id = f"MS-{len(milestones)+1:02d}"
            milestones.append({"id": ms_id, "name": name, "target_date": date,
                                "status": status or "Planned", "owner": owner})
        if milestones:
            result["milestones"] = milestones
        break

    # ---- Tech Stack table
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if any(k in combined for k in ["frontend", "backend", "database", "infrastructure",
                                        "tech", "stack", "เทคโนโลยี"]):
            ts = existing.get("tech_stack", {}).copy()
            ts["frontend"]       = kv_lookup(kv, "frontend")  or ts.get("frontend", "")
            ts["backend"]        = kv_lookup(kv, "backend")   or ts.get("backend", "")
            ts["database"]       = kv_lookup(kv, "database", "db", "ฐานข้อมูล") or ts.get("database", "")
            ts["infrastructure"] = kv_lookup(kv, "infrastructure", "infra", "cloud", "server") or ts.get("infrastructure", "")
            ts["source_control"] = kv_lookup(kv, "source control", "git", "source") or ts.get("source_control", "")
            ts["ci_cd"]          = kv_lookup(kv, "ci/cd", "ci_cd", "pipeline", "deploy") or ts.get("ci_cd", "")
            result["tech_stack"] = ts
            break

    return result


def extract_mm(doc, existing: dict) -> dict:
    """Extract from Meeting Minutes → meetings."""
    meetings = []
    # Use find_marker_blocks with MTG-xxx pattern (documents use [normal] style with MTG id in heading)
    blocks = find_marker_blocks(doc, r'MTG-\d+')

    for heading, tables in blocks:
        mtg_id_match = re.search(r"MTG-\d+", heading, re.IGNORECASE)
        if not mtg_id_match:
            continue
        mtg_id = mtg_id_match.group(0).upper()

        title = re.sub(r"บันทึกการประชุม\s*[—–-]\s*", "", heading, flags=re.IGNORECASE).strip()
        title = re.sub(r"\(\s*MTG-\d+\s*\)", "", title, flags=re.IGNORECASE).strip()
        title = re.sub(r"\s*[-—]\s*$", "", title).strip()
        if not title:
            title = heading.strip()

        # Find meeting info KV table (has Date/Time/Location/Chair keys)
        kv = {}
        for tbl in tables:
            kv_candidate = read_kv_table(tbl)
            combined = " ".join(kv_candidate.keys()).lower()
            if any(k in combined for k in ["วันที่", "date", "เวลา", "time", "สถานที่", "location", "chair", "ประธาน", "mtg", "meeting id"]):
                kv = kv_candidate
                break
        date_    = kv_lookup(kv, "date", "วันที่") or ""
        time_    = kv_lookup(kv, "time", "เวลา") or ""
        location = kv_lookup(kv, "location", "สถานที่") or ""
        chair    = kv_lookup(kv, "chair", "ประธาน") or ""

        # Attendees table (has #/Name/Signature columns)
        attendees = []
        for tbl in tables:
            headers_lower = [_cell_text(c).lower() for c in tbl.rows[0].cells] if tbl.rows else []
            if any("name" in h or "ชื่อ" in h for h in headers_lower) and len(headers_lower) >= 3:
                for row in tbl.rows[1:]:
                    cells = [_cell_text(c) for c in row.cells]
                    name_idx = next((i for i, h in enumerate(headers_lower) if "name" in h or "ชื่อ" in h), 1)
                    if name_idx < len(cells) and cells[name_idx]:
                        attendees.append(cells[name_idx])
                break

        # Action items table
        action_items = []
        for tbl in tables:
            headers_lower = [_cell_text(c).lower() for c in tbl.rows[0].cells] if tbl.rows else []
            if any("action" in h or "รายการ" in h for h in headers_lower) and \
               any("owner" in h or "due" in h or "ผู้รับผิดชอบ" in h for h in headers_lower):
                rows = read_row_table(tbl)
                for row in rows:
                    item   = row.get("รายการ / Action", row.get("Action", row.get("รายการ", ""))).strip()
                    owner  = row.get("ผู้รับผิดชอบ / Owner", row.get("Owner", "")).strip()
                    due    = row.get("กำหนด / Due", row.get("Due", row.get("กำหนด", ""))).strip()
                    status = row.get("สถานะ / Status", row.get("Status", "Open")).strip()
                    if item:
                        action_items.append({"item": item, "owner": owner, "due_date": due, "status": status or "Open"})
                break

        meetings.append({
            "id":           mtg_id,
            "title":        title,
            "date":         date_,
            "time":         time_,
            "location":     location,
            "chair":        chair,
            "attendees":    attendees,
            "agenda":       [],
            "action_items": action_items,
            "summary":      "",
        })

    return {"meetings": meetings} if meetings else {}


def extract_sow(doc, existing: dict) -> dict:
    """Extract from Statement of Work → sow."""
    sow = existing.get("sow", {}).copy()

    # Contract overview table
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if any(k in combined for k in ["contract", "สัญญา", "effective", "expiration"]):
            sow["contract_number"] = kv_lookup(kv, "contract number", "เลขที่สัญญา", "contract no") or sow.get("contract_number", "")
            sow["effective_date"]  = kv_lookup(kv, "effective date", "วันที่มีผล", "start") or sow.get("effective_date", "")
            sow["expiration_date"] = kv_lookup(kv, "expiration", "expiry", "end date", "วันสิ้นสุด") or sow.get("expiration_date", "")
            break

    # Client / Vendor tables
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if "client" in combined or "ลูกค้า" in combined:
            client = sow.get("client", {})
            client["name"]           = kv_lookup(kv, "organization", "หน่วยงาน", "name") or client.get("name", "")
            client["representative"] = kv_lookup(kv, "representative", "ผู้แทน") or client.get("representative", "")
            client["title"]          = kv_lookup(kv, "title", "ตำแหน่ง") or client.get("title", "")
            sow["client"] = client
        elif "vendor" in combined or "ผู้พัฒนา" in combined:
            vendor = sow.get("vendor", {})
            vendor["name"]           = kv_lookup(kv, "organization", "บริษัท", "vendor") or vendor.get("name", "")
            vendor["representative"] = kv_lookup(kv, "representative", "ผู้แทน") or vendor.get("representative", "")
            vendor["title"]          = kv_lookup(kv, "title", "ตำแหน่ง") or vendor.get("title", "")
            sow["vendor"] = vendor

    # Deliverables table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("deliverable" in h or "del-" in h or "งานที่ส่งมอบ" in h for h in headers_lower):
            rows = read_row_table(table)
            deliverables = []
            for row in rows:
                del_id      = row.get("ID", row.get("DEL ID", "")).strip()
                description = row.get("Description", row.get("คำอธิบาย", row.get("Deliverable", ""))).strip()
                acceptance  = row.get("Acceptance Criteria", row.get("เกณฑ์ยอมรับ", "")).strip()
                due_date    = row.get("Due Date", row.get("Due", row.get("วันส่งมอบ", ""))).strip()
                milestone   = row.get("Payment Milestone", row.get("Milestone", "")).strip()
                if description:
                    if not del_id:
                        del_id = f"DEL-{len(deliverables)+1:03d}"
                    deliverables.append({"id": del_id, "description": description,
                                         "acceptance_criteria": acceptance,
                                         "due_date": due_date, "payment_milestone": milestone})
            if deliverables:
                sow["deliverables"] = deliverables
            break

    # Payment schedule table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("payment" in h or "percentage" in h or "%" in h or "ชำระ" in h for h in headers_lower):
            rows = read_row_table(table)
            schedule = []
            for row in rows:
                ms      = row.get("Milestone", row.get("Milestone Payment", row.get("เหตุการณ์", ""))).strip()
                pct_raw = row.get("Percentage", row.get("%", row.get("สัดส่วน", ""))).strip()
                amount  = row.get("Amount", row.get("จำนวนเงิน", "")).strip()
                due     = row.get("Due Date", row.get("Due", row.get("กำหนด", ""))).strip()
                if ms:
                    try:
                        pct = int(re.sub(r"[^0-9]", "", pct_raw)) if pct_raw else 0
                    except ValueError:
                        pct = 0
                    schedule.append({"milestone": ms, "percentage": pct, "amount": amount, "due_date": due})
            if schedule:
                sow["payment_schedule"] = schedule
            break

    return {"sow": sow} if sow else {}


def extract_brd(doc, existing: dict) -> dict:
    """Extract from Requirements Document → requirements."""
    reqs = []
    existing_reqs = {r.get("id", ""): r for r in existing.get("requirements", [])}
    # Use find_marker_blocks: documents use [normal] paragraphs with ▌ REQ-xxx pattern
    blocks = find_marker_blocks(doc, r'\bREQ-\d+')

    for heading, tables in blocks:
        req_id = extract_id_from_heading(heading, "REQ")
        if not req_id.startswith("REQ-"):
            continue
        title = extract_title_from_heading(heading, req_id)
        ex = existing_reqs.get(req_id, {})

        kv = read_kv_table(tables[0]) if tables else {}
        req = {
            "id":                  req_id,
            "title":               title,
            "description":         kv_lookup(kv, "description", "คำอธิบาย") or ex.get("description", ""),
            "priority":            kv_lookup(kv, "priority", "ลำดับความสำคัญ") or ex.get("priority", "Medium"),
            "type":                kv_lookup(kv, "type", "ประเภท") or ex.get("type", "Functional"),
            "category":            kv_lookup(kv, "category", "หมวดหมู่") or ex.get("category", ""),
            "source":              kv_lookup(kv, "source", "แหล่งที่มา") or ex.get("source", ""),
            "acceptance_criteria": kv_lookup(kv, "acceptance", "เกณฑ์ยอมรับ") or ex.get("acceptance_criteria", ""),
            "linked_design":       ex.get("linked_design", []),
            "linked_test_cases":   ex.get("linked_test_cases", []),
        }
        reqs.append(req)

    return {"requirements": reqs} if reqs else {}


def extract_sdd(doc, existing: dict) -> dict:
    """Extract from System Design Document → design_components."""
    comps = []
    existing_comps = {c.get("id", ""): c for c in existing.get("design_components", [])}

    # Look for the Components table (COMP ID, Name, Type, Description, REQ Reference)
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("comp" in h for h in headers_lower) and any("req" in h for h in headers_lower):
            rows = read_row_table(table)
            for row in rows:
                comp_id = row.get("COMP ID", row.get("ID", "")).strip()
                if not comp_id or not comp_id.startswith("COMP"):
                    continue
                ex = existing_comps.get(comp_id, {})
                name     = row.get("ชื่อ", row.get("Name", "")).strip()
                type_    = row.get("ประเภท", row.get("Type", "")).strip()
                desc     = row.get("คำอธิบาย", row.get("Description", "")).strip()
                req_ref  = row.get("REQ Reference", row.get("REQ Ref", "")).strip()
                req_list = [r.strip() for r in re.split(r"[,\s]+", req_ref) if r.strip().startswith("REQ-")]
                comps.append({
                    "id": comp_id,
                    "name": name or ex.get("name", ""),
                    "description": desc or ex.get("description", ""),
                    "type": type_ or ex.get("type", ""),
                    "related_requirements": req_list or ex.get("related_requirements", []),
                    "technology": ex.get("technology", ""),
                })
            break

    return {"design_components": comps} if comps else {}


def extract_dbd(doc, existing: dict) -> dict:
    """Extract from Database Design Document → database_tables."""
    tables_data = []
    existing_tables = {t.get("name", ""): t for t in existing.get("database_tables", [])}

    # Strategy: use overview table (Table Name, Description, REQ Reference) for names/reqs
    # and column detail tables (Column Name, Data Type, Description) for column structure.
    # Also use find_marker_blocks to get section headings for table names.

    # 1. Get table names/descriptions from section headings (▌ Table: xxx — description)
    blocks = find_marker_blocks(doc, r'Table\s*:\s*\w+')
    section_info = []  # [(table_name, description), ...]
    for heading, _ in blocks:
        m = re.match(r".*?Table\s*:\s*(\w+)\s*[—–-]?\s*(.*)", heading, re.IGNORECASE)
        if m:
            section_info.append((m.group(1).strip(), m.group(2).strip()))

    # 2. Get overview table for REQ references
    req_refs_by_name = {}
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if "table name" in " ".join(headers_lower) and any("req" in h for h in headers_lower):
            for row in read_row_table(table):
                tname = row.get("Table Name", "").strip()
                req_ref = row.get("REQ Reference", row.get("REQ Ref", "")).strip()
                if tname:
                    req_refs_by_name[tname] = [r.strip() for r in re.split(r"[,\s]+", req_ref)
                                                if r.strip().startswith("REQ-")]
            break

    # 3. Collect all column detail tables (Column Name, Data Type, Description)
    col_tables = []
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("column name" in h or ("column" in h and "name" in h) for h in headers_lower) and \
           any("data type" in h or "type" in h for h in headers_lower):
            col_tables.append(table)

    # 4. Match column tables with section names
    for i, col_table in enumerate(col_tables):
        if i < len(section_info):
            tbl_name, tbl_desc = section_info[i]
        else:
            tbl_name = f"Table_{i+1}"
            tbl_desc = ""
        ex = existing_tables.get(tbl_name, {})
        req_list = req_refs_by_name.get(tbl_name, ex.get("related_requirements", []))
        col_rows = read_row_table(col_table)
        columns = []
        for row in col_rows:
            col_name = row.get("Column Name", row.get("Name", "")).strip()
            col_type = row.get("Data Type", row.get("Type", "")).strip()
            col_desc = row.get("Description", row.get("คำอธิบาย", "")).strip()
            if col_name:
                columns.append({"name": col_name, "type": col_type, "description": col_desc})
        tables_data.append({
            "name": tbl_name,
            "description": tbl_desc or ex.get("description", ""),
            "columns": columns or ex.get("columns", []),
            "related_requirements": req_list,
        })

    return {"database_tables": tables_data} if tables_data else {}


def extract_tcr(doc, existing: dict) -> dict:
    """Extract from Test Cases & Results → test_cases."""
    tcs = []
    existing_tcs = {tc.get("id", ""): tc for tc in existing.get("test_cases", [])}
    # Use find_marker_blocks: documents use [normal] paragraphs with ▌ TC-xxx pattern
    blocks = find_marker_blocks(doc, r'\bTC-\d+')

    for heading, tables in blocks:
        tc_id = extract_id_from_heading(heading, "TC")
        if not tc_id.startswith("TC-"):
            continue
        title = extract_title_from_heading(heading, tc_id)
        ex = existing_tcs.get(tc_id, {})
        kv = read_kv_table(tables[0]) if tables else {}

        steps_raw = kv_lookup(kv, "test steps", "steps") or ex.get("steps", [])
        if isinstance(steps_raw, str):
            steps = [s.strip() for s in steps_raw.split("\n") if s.strip()] or [steps_raw]
        else:
            steps = steps_raw

        tc = {
            "id":                  tc_id,
            "title":               title,
            "related_requirement": kv_lookup(kv, "req reference", "req ref", "requirement") or ex.get("related_requirement", ""),
            "test_type":           kv_lookup(kv, "test type") or ex.get("test_type", "Functional"),
            "preconditions":       kv_lookup(kv, "precondition") or ex.get("preconditions", ""),
            "steps":               steps,
            "expected_result":     kv_lookup(kv, "expected result", "expected") or ex.get("expected_result", ""),
            "actual_result":       kv_lookup(kv, "actual result", "actual") or ex.get("actual_result", ""),
            "status":              kv_lookup(kv, "status", "สถานะ") or ex.get("status", "Pending"),
            "tester":              kv_lookup(kv, "tester") or ex.get("tester", ""),
            "test_date":           kv_lookup(kv, "test date", "date") or ex.get("test_date", ""),
            "remarks":             kv_lookup(kv, "remarks", "หมายเหตุ") or ex.get("remarks", ""),
        }
        tcs.append(tc)

    return {"test_cases": tcs} if tcs else {}


def extract_bdl(doc, existing: dict) -> dict:
    """Extract from Bug/Defect Log → defects."""
    defects = []
    existing_defects = {d.get("id", ""): d for d in existing.get("defects", [])}

    # Summary list table (BUG ID, Title, Severity, Status, REQ, TC, ...)
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("bug" in h or "id" in h for h in headers_lower) and \
           any("severity" in h for h in headers_lower):
            rows = read_row_table(table)
            for row in rows:
                bug_id = row.get("ID", row.get("BUG ID", row.get("Bug ID", ""))).strip()
                if not bug_id.startswith("BUG"):
                    continue
                ex = existing_defects.get(bug_id, {})
                title    = row.get("Title", "").strip()
                severity = row.get("Severity", "").strip()
                priority = row.get("Priority", "").strip()
                status   = row.get("Status", "").strip()
                req_ref  = row.get("REQ", row.get("REQ Ref", "")).strip()
                tc_ref   = row.get("TC", row.get("TC Ref", "")).strip()
                reporter = row.get("Reported By", "").strip()
                rep_date = row.get("Date", row.get("Reported Date", "")).strip()
                assigned = row.get("Assigned", row.get("Assigned To", "")).strip()
                fixed_d  = row.get("Fixed Date", "").strip()
                defects.append({
                    "id":                bug_id,
                    "title":             title or ex.get("title", ""),
                    "related_test_case": tc_ref or ex.get("related_test_case", ""),
                    "related_requirement": req_ref or ex.get("related_requirement", ""),
                    "severity":          severity or ex.get("severity", "Medium"),
                    "priority":          priority or ex.get("priority", "Medium"),
                    "status":            status or ex.get("status", "Open"),
                    "reported_by":       reporter or ex.get("reported_by", ""),
                    "reported_date":     rep_date or ex.get("reported_date", ""),
                    "assigned_to":       assigned or ex.get("assigned_to", ""),
                    "fixed_date":        fixed_d or ex.get("fixed_date", ""),
                    "description":       ex.get("description", ""),
                    "steps_to_reproduce": ex.get("steps_to_reproduce", ""),
                    "root_cause":        ex.get("root_cause", ""),
                    "resolution":        ex.get("resolution", ""),
                })
            break

    # Detail blocks for description/root_cause/resolution
    blocks = find_marker_blocks(doc, r'\bBUG-\d+')
    defect_map = {d["id"]: d for d in defects}
    for heading, tables in blocks:
        bug_id = extract_id_from_heading(heading, "BUG")
        if not bug_id.startswith("BUG-"):
            continue
        kv = read_kv_table(tables[0]) if tables else {}
        if bug_id in defect_map:
            defect_map[bug_id]["description"]        = kv_lookup(kv, "description", "คำอธิบาย") or defect_map[bug_id].get("description", "")
            defect_map[bug_id]["steps_to_reproduce"] = kv_lookup(kv, "steps to reproduce", "steps") or defect_map[bug_id].get("steps_to_reproduce", "")
            defect_map[bug_id]["root_cause"]         = kv_lookup(kv, "root cause") or defect_map[bug_id].get("root_cause", "")
            defect_map[bug_id]["resolution"]         = kv_lookup(kv, "resolution") or defect_map[bug_id].get("resolution", "")

    return {"defects": defects} if defects else {}


def extract_rr(doc, existing: dict) -> dict:
    """Extract from Risk Register → risks."""
    risks = []
    existing_risks = {r.get("id", ""): r for r in existing.get("risks", [])}

    # Summary table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("risk id" in h or "risk" in h for h in headers_lower) and \
           any("probability" in h or "prob" in h or "impact" in h or "level" in h for h in headers_lower):
            rows = read_row_table(table)
            for row in rows:
                risk_id = row.get("RISK ID", row.get("Risk ID", row.get("ID", ""))).strip()
                if not risk_id.startswith("RISK"):
                    continue
                ex = existing_risks.get(risk_id, {})
                risks.append({
                    "id":          risk_id,
                    "category":    row.get("Category", "").strip() or ex.get("category", ""),
                    "description": row.get("Description", "").strip() or ex.get("description", ""),
                    "probability": row.get("Prob.", row.get("Probability", "")).strip() or ex.get("probability", ""),
                    "impact":      row.get("Impact", "").strip() or ex.get("impact", ""),
                    "risk_level":  row.get("Level", row.get("Risk Level", "")).strip() or ex.get("risk_level", ""),
                    "mitigation":  row.get("Mitigation", "").strip() or ex.get("mitigation", ""),
                    "owner":       row.get("Owner", "").strip() or ex.get("owner", ""),
                    "status":      row.get("Status", "").strip() or ex.get("status", "Open"),
                    "review_date": row.get("Review Date", "").strip() or ex.get("review_date", ""),
                    "contingency": ex.get("contingency", ""),
                    "linked_capa": ex.get("linked_capa", ""),
                })
            break

    # Detail blocks to enrich contingency/linked_capa
    blocks = find_marker_blocks(doc, r'\bRISK-\d+')
    risk_map = {r["id"]: r for r in risks}
    for heading, tables in blocks:
        risk_id = extract_id_from_heading(heading, "RISK")
        if not risk_id.startswith("RISK-"):
            continue
        kv = read_kv_table(tables[0]) if tables else {}
        ex = existing_risks.get(risk_id, {})
        if risk_id in risk_map:
            risk_map[risk_id]["contingency"] = kv_lookup(kv, "contingency plan", "contingency") or ex.get("contingency", "")
            risk_map[risk_id]["linked_capa"] = kv_lookup(kv, "linked capa", "capa") or ex.get("linked_capa", "")
            risk_map[risk_id]["mitigation"]  = kv_lookup(kv, "mitigation plan", "mitigation") or risk_map[risk_id].get("mitigation", "")

    return {"risks": risks} if risks else {}


def extract_crf(doc, existing: dict) -> dict:
    """Extract from Change Request Form → change_requests."""
    changes = []
    existing_changes = {cr.get("id", ""): cr for cr in existing.get("change_requests", [])}
    # Use find_marker_blocks: documents use [normal] paragraphs like "Change Request: CR-001 —"
    blocks = find_marker_blocks(doc, r'(?i)change request.*?CR-\d+|\bCR-\d+')

    for heading, tables in blocks:
        if not any(kw in heading.lower() for kw in ["change request", "cr-", "crf-"]):
            continue
        cr_id_m = re.search(r"(CRF?-\d+)", heading, re.IGNORECASE)
        cr_id = cr_id_m.group(1).upper() if cr_id_m else ""
        ex = existing_changes.get(cr_id, {})

        # Combine all KV tables in this block
        combined_kv = {}
        for tbl in tables:
            combined_kv.update(read_kv_table(tbl))

        cr = {
            "id":                   cr_id or kv_lookup(combined_kv, "cr id") or ex.get("id", ""),
            "title":                kv_lookup(combined_kv, "title", "ชื่อ") or ex.get("title", ""),
            "description":          ex.get("description", ""),
            "requestor":            kv_lookup(combined_kv, "requestor", "ผู้ขอ") or ex.get("requestor", ""),
            "request_date":         kv_lookup(combined_kv, "request date", "วันที่ขอ") or ex.get("request_date", ""),
            "priority":             kv_lookup(combined_kv, "priority", "ลำดับ") or ex.get("priority", "Medium"),
            "impact":               ex.get("impact", ""),
            "affected_documents":   ex.get("affected_documents", []),
            "status":               kv_lookup(combined_kv, "status", "สถานะ") or ex.get("status", "Pending"),
            "approved_by":          kv_lookup(combined_kv, "approved by", "อนุมัติโดย") or ex.get("approved_by", ""),
            "approval_date":        kv_lookup(combined_kv, "approval date", "วันที่อนุมัติ") or ex.get("approval_date", ""),
            "implementation_date":  kv_lookup(combined_kv, "implementation date", "วันที่ดำเนินการ") or ex.get("implementation_date", ""),
        }
        if cr["id"]:
            changes.append(cr)

    return {"change_requests": changes} if changes else {}


def extract_vrn(doc, existing: dict) -> dict:
    """Extract from Version Release Notes → versions + deployments."""
    versions = []
    deployments = []
    existing_versions = {v.get("version", ""): v for v in existing.get("versions", [])}
    existing_deps = {d.get("id", ""): d for d in existing.get("deployments", [])}

    # First pass: extract from version summary table (Version, Release Date, Type, Description)
    version_from_summary = {}
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if "version" in headers_lower and any("release" in h or "type" in h for h in headers_lower):
            for row in read_row_table(table):
                ver = row.get("Version", "").strip()
                if ver and re.match(r'\d+\.\d+', ver):
                    version_from_summary[ver] = {
                        "release_date": row.get("Release Date", row.get("Date", "")).strip(),
                        "release_type": row.get("Type", row.get("Release Type", "Major")).strip(),
                        "description":  row.get("Description", "").strip(),
                    }
            break

    # Second pass: detailed blocks per Release (▌ or "Release v1.0.0 —" paragraphs)
    blocks = find_marker_blocks(doc, r'(?i)release\s+v\d+\.\d+')
    for heading, tables in blocks:
        m = re.search(r'v(\d+[\d.]+)', heading, re.IGNORECASE)
        if not m:
            continue
        ver_str = m.group(1).rstrip('.')
        ex = existing_versions.get(ver_str, {})
        summary = version_from_summary.get(ver_str, {})
        combined_kv = {}
        for tbl in tables:
            combined_kv.update(read_kv_table(tbl))

        ver = {
            "version":      ver_str,
            "release_date": kv_lookup(combined_kv, "release date", "date") or summary.get("release_date") or ex.get("release_date", ""),
            "release_type": kv_lookup(combined_kv, "release type", "type") or summary.get("release_type") or ex.get("release_type", "Major"),
            "description":  kv_lookup(combined_kv, "description") or summary.get("description") or ex.get("description", ""),
            "changes":      ex.get("changes", []),
            "deployed_by":  kv_lookup(combined_kv, "deployed by", "deployer") or ex.get("deployed_by", ""),
            "environment":  kv_lookup(combined_kv, "environment", "env") or ex.get("environment", "Production"),
        }
        versions.append(ver)
        version_from_summary.pop(ver_str, None)  # mark as handled

        dep_id_raw   = kv_lookup(combined_kv, "deployment id", "dep id", "deploy id") or ""
        dep_status   = kv_lookup(combined_kv, "status") or ""
        dep_approval = kv_lookup(combined_kv, "approved by", "approval") or ""
        dep_type     = kv_lookup(combined_kv, "deployment type", "type") or ""
        if dep_id_raw and dep_id_raw.startswith("DEP"):
            ex_dep = existing_deps.get(dep_id_raw, {})
            deployments.append({
                "id":              dep_id_raw,
                "version":         ver_str,
                "date":            kv_lookup(combined_kv, "date") or ex_dep.get("date", ""),
                "environment":     kv_lookup(combined_kv, "environment") or ex_dep.get("environment", "Production"),
                "deployed_by":     ver["deployed_by"],
                "deployment_type": dep_type or ex_dep.get("deployment_type", ""),
                "steps":           ex_dep.get("steps", []),
                "rollback_plan":   ex_dep.get("rollback_plan", ""),
                "status":          dep_status or ex_dep.get("status", ""),
                "approval":        dep_approval or ex_dep.get("approval", ""),
            })

    # Fallback: use summary table data for any version not found in blocks
    for ver_str, summary in version_from_summary.items():
        ex = existing_versions.get(ver_str, {})
        versions.append({
            "version":      ver_str,
            "release_date": summary.get("release_date", ""),
            "release_type": summary.get("release_type", "Major"),
            "description":  summary.get("description", ""),
            "changes":      ex.get("changes", []),
            "deployed_by":  ex.get("deployed_by", ""),
            "environment":  ex.get("environment", "Production"),
        })

    result = {}
    if versions:
        result["versions"] = versions
    if deployments:
        result["deployments"] = deployments
    return result


def extract_tr(doc, existing: dict) -> dict:
    """Extract from Training Record → training_sessions."""
    sessions = []
    existing_sessions = {s.get("id", ""): s for s in existing.get("training_sessions", [])}
    # Documents use [normal] paragraphs: "Training Session: Project Name (TRN-001)"
    blocks = find_marker_blocks(doc, r'(?i)training session\s*:|\bTRN-\d+')

    for heading, tables in blocks:
        if not any(kw in heading.lower() for kw in ["training session", "trn-", "การฝึกอบรม"]):
            continue
        trn_id_m = re.search(r"TRN-\d+", heading, re.IGNORECASE)
        trn_id = trn_id_m.group(0).upper() if trn_id_m else ""
        title_raw = re.sub(r"Training Session\s*:\s*", "", heading, flags=re.IGNORECASE)
        title_raw = re.sub(r"\(TRN-\d+\)", "", title_raw, flags=re.IGNORECASE).strip()
        ex = existing_sessions.get(trn_id, {})
        combined_kv = {}
        for tbl in tables:
            combined_kv.update(read_kv_table(tbl))

        # Attendees
        attendees = []
        for tbl in tables:
            headers_lower = [_cell_text(c).lower() for c in tbl.rows[0].cells] if tbl.rows else []
            if any("name" in h or "ชื่อ" in h for h in headers_lower) and \
               any("dept" in h or "หน่วยงาน" in h or "signature" in h or "ลายมือ" in h for h in headers_lower):
                for row in tbl.rows[1:]:
                    cells = [_cell_text(c) for c in row.cells]
                    name_idx = next((i for i, h in enumerate(headers_lower) if "name" in h or "ชื่อ" in h), 1)
                    dept_idx = next((i for i, h in enumerate(headers_lower) if "dept" in h or "หน่วยงาน" in h), 2)
                    sig_idx  = next((i for i, h in enumerate(headers_lower) if "sign" in h or "ลายมือ" in h), 3)
                    name = cells[name_idx] if name_idx < len(cells) else ""
                    dept = cells[dept_idx] if dept_idx < len(cells) else ""
                    sig  = cells[sig_idx]  if sig_idx  < len(cells) else ""
                    if name:
                        attendees.append({"name": name, "department": dept, "signed": bool(sig and sig not in ["☐", "[ ]", ""])})
                break

        sessions.append({
            "id":       trn_id or f"TRN-{len(sessions)+1:03d}",
            "title":    title_raw or ex.get("title", ""),
            "date":     kv_lookup(combined_kv, "date", "วันที่") or ex.get("date", ""),
            "duration": kv_lookup(combined_kv, "duration", "ระยะเวลา") or ex.get("duration", ""),
            "location": kv_lookup(combined_kv, "location", "สถานที่") or ex.get("location", ""),
            "trainer":  kv_lookup(combined_kv, "trainer", "วิทยากร") or ex.get("trainer", ""),
            "topics":   ex.get("topics", []),
            "attendees": attendees or ex.get("attendees", []),
        })

    return {"training_sessions": sessions} if sessions else {}


def extract_isl(doc, existing: dict) -> dict:
    """Extract from Incident/Support Log → incidents."""
    incidents = []
    existing_incidents = {i.get("id", ""): i for i in existing.get("incidents", [])}

    # Summary table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("inc" in h for h in headers_lower) and any("severity" in h for h in headers_lower):
            rows = read_row_table(table)
            for row in rows:
                inc_id = row.get("ID", row.get("INC ID", "")).strip()
                if not inc_id.startswith("INC"):
                    continue
                ex = existing_incidents.get(inc_id, {})
                incidents.append({
                    "id":                     inc_id,
                    "title":                  row.get("Title", "").strip() or ex.get("title", ""),
                    "date_reported":          row.get("Date", row.get("Date Reported", "")).strip() or ex.get("date_reported", ""),
                    "severity":               row.get("Severity", "").strip() or ex.get("severity", ""),
                    "description":            ex.get("description", ""),
                    "reported_by":            ex.get("reported_by", ""),
                    "assigned_to":            row.get("Assigned To", row.get("Assigned", "")).strip() or ex.get("assigned_to", ""),
                    "status":                 row.get("Status", "").strip() or ex.get("status", "Open"),
                    "resolution":             ex.get("resolution", ""),
                    "resolved_date":          row.get("Resolved", row.get("Resolved Date", "")).strip() or ex.get("resolved_date", ""),
                    "root_cause":             ex.get("root_cause", ""),
                    "linked_change_request":  row.get("CR Link", "").strip() or ex.get("linked_change_request", ""),
                })
            break

    # Detail blocks
    blocks = find_marker_blocks(doc, r'\bINC-\d+')
    inc_map = {i["id"]: i for i in incidents}
    for heading, tables in blocks:
        inc_id = extract_id_from_heading(heading, "INC")
        if not inc_id.startswith("INC-"):
            continue
        kv = read_kv_table(tables[0]) if tables else {}
        ex = existing_incidents.get(inc_id, {})
        if inc_id in inc_map:
            inc_map[inc_id]["description"]  = kv_lookup(kv, "description", "คำอธิบาย") or ex.get("description", "")
            inc_map[inc_id]["root_cause"]   = kv_lookup(kv, "root cause") or ex.get("root_cause", "")
            inc_map[inc_id]["resolution"]   = kv_lookup(kv, "resolution") or ex.get("resolution", "")
        else:
            # New incident found only in detail block
            title_from_heading = extract_title_from_heading(heading, inc_id)
            incidents.append({
                "id":                    inc_id,
                "title":                 title_from_heading,
                "date_reported":         ex.get("date_reported", ""),
                "severity":              ex.get("severity", ""),
                "description":           kv_lookup(kv, "description") or ex.get("description", ""),
                "reported_by":           ex.get("reported_by", ""),
                "assigned_to":           ex.get("assigned_to", ""),
                "status":                ex.get("status", "Open"),
                "resolution":            kv_lookup(kv, "resolution") or ex.get("resolution", ""),
                "resolved_date":         ex.get("resolved_date", ""),
                "root_cause":            kv_lookup(kv, "root cause") or ex.get("root_cause", ""),
                "linked_change_request": ex.get("linked_change_request", ""),
            })

    return {"incidents": incidents} if incidents else {}


def extract_ar(doc, existing: dict) -> dict:
    """Extract from Audit Report → audit."""
    audit = existing.get("audit", {}).copy()

    # Audit information KV table
    for table in doc.tables:
        kv = read_kv_table(table)
        combined = " ".join(kv.keys()).lower()
        if any(k in combined for k in ["audit date", "auditor", "audit scope"]):
            audit["audit_date"]  = kv_lookup(kv, "audit date", "date") or audit.get("audit_date", "")
            audit["auditor"]     = kv_lookup(kv, "auditor") or audit.get("auditor", "")
            audit["audit_scope"] = kv_lookup(kv, "audit scope", "scope") or audit.get("audit_scope", "")
            break

    # Findings table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("finding" in h for h in headers_lower) and any("clause" in h for h in headers_lower):
            rows = read_row_table(table)
            findings = []
            for row in rows:
                fid    = row.get("Finding ID", row.get("ID", "")).strip()
                clause = row.get("ISO Clause", row.get("Clause", "")).strip()
                ftype  = row.get("Type", row.get("Finding Type", "")).strip()
                desc   = row.get("Description", "").strip()
                capa   = row.get("Linked CAPA", row.get("CAPA", "")).strip()
                if fid and (clause or desc):
                    findings.append({"id": fid, "clause": clause, "finding_type": ftype,
                                     "description": desc, "linked_capa": capa})
            if findings:
                audit["findings"] = findings
            break

    return {"audit": audit} if audit else {}


def extract_capa(doc, existing: dict) -> dict:
    """Extract from CAPA document → capas."""
    capas = []
    existing_capas = {c.get("id", ""): c for c in existing.get("capas", [])}

    # CAPA Register table
    for table in doc.tables:
        headers_lower = [_cell_text(c).lower() for c in table.rows[0].cells] if table.rows else []
        if any("capa" in h for h in headers_lower) and any("type" in h for h in headers_lower):
            rows = read_row_table(table)
            for row in rows:
                capa_id = row.get("CAPA ID", row.get("ID", "")).strip()
                if not capa_id.startswith("CAPA"):
                    continue
                ex = existing_capas.get(capa_id, {})
                capas.append({
                    "id":                  capa_id,
                    "type":                row.get("Type", "").strip() or ex.get("type", ""),
                    "related_finding":     row.get("Finding/Risk", row.get("Finding", "")).strip() or ex.get("related_finding", ""),
                    "description":         row.get("Description", "").strip() or ex.get("description", ""),
                    "root_cause":          ex.get("root_cause", ""),
                    "action_plan":         ex.get("action_plan", ""),
                    "responsible":         row.get("Owner", "").strip() or ex.get("responsible", ""),
                    "target_date":         row.get("Target Date", "").strip() or ex.get("target_date", ""),
                    "status":              row.get("Status", "").strip() or ex.get("status", "Open"),
                    "effectiveness_review": ex.get("effectiveness_review", ""),
                    "closed_date":         row.get("Closed Date", "").strip() or ex.get("closed_date", ""),
                })
            break

    # Detail blocks to enrich root_cause, action_plan, effectiveness_review
    blocks = find_marker_blocks(doc, r'\bCAPA-\d+')
    capa_map = {c["id"]: c for c in capas}
    for heading, tables in blocks:
        capa_id = extract_id_from_heading(heading, "CAPA")
        if not capa_id.startswith("CAPA-"):
            continue
        kv = read_kv_table(tables[0]) if tables else {}
        ex = existing_capas.get(capa_id, {})
        if capa_id in capa_map:
            capa_map[capa_id]["root_cause"]           = kv_lookup(kv, "root cause") or ex.get("root_cause", "")
            capa_map[capa_id]["action_plan"]          = kv_lookup(kv, "action plan") or ex.get("action_plan", "")
            capa_map[capa_id]["effectiveness_review"] = kv_lookup(kv, "effectiveness review", "effectiveness") or ex.get("effectiveness_review", "")

    return {"capas": capas} if capas else {}


# ─── Document file finder ─────────────────────────────────────────────────────

def find_project_docs(project_path: Path) -> dict:
    """
    Scan project folders for .docx files.
    Returns dict: {doc_code: Path, ...}
    Priority: if multiple files match a code, pick the most recently modified.
    """
    doc_map = {}
    for folder_name, codes in FOLDER_DOC_MAP.items():
        folder = project_path / folder_name
        if not folder.exists():
            continue
        for docx_file in sorted(folder.glob("*.docx")):
            stem = docx_file.stem.upper()
            for code in codes:
                # Match e.g. "DVETS-01-PP-v1.0" contains "-PP-"
                if f"-{code}-" in stem or stem.endswith(f"-{code}") or stem.startswith(f"{code}-"):
                    existing = doc_map.get(code)
                    if existing is None or docx_file.stat().st_mtime > existing.stat().st_mtime:
                        doc_map[code] = docx_file
                    break
    return doc_map


# ─── Main extraction orchestrator ─────────────────────────────────────────────

EXTRACTOR_MAP = {
    "PP":   extract_pp,
    "MM":   extract_mm,
    "SOW":  extract_sow,
    "BRD":  extract_brd,
    "SDD":  extract_sdd,
    "DBD":  extract_dbd,
    "TCR":  extract_tcr,
    "BDL":  extract_bdl,
    "RR":   extract_rr,
    "CRF":  extract_crf,
    "VRN":  extract_vrn,
    "TR":   extract_tr,
    "ISL":  extract_isl,
    "AR":   extract_ar,
    "CAPA": extract_capa,
}


def extract_from_project(project_path: Path, existing_config: dict, verbose: bool = False) -> dict:
    """
    Scan all documents and merge extracted data into a new config dict.
    """
    config = existing_config.copy()
    doc_map = find_project_docs(project_path)

    if not doc_map:
        print(f"  WARNING: No .docx files found in {project_path}")
        return config

    print(f"  Found {len(doc_map)} documents to process: {', '.join(sorted(doc_map.keys()))}")

    for code in ["PP", "MM", "SOW", "BRD", "SDD", "DBD", "TCR", "BDL",
                 "RR", "CRF", "VRN", "TR", "ISL", "AR", "CAPA"]:
        if code not in doc_map:
            if verbose:
                print(f"    [{code:4s}] — not found, skipping")
            continue

        extractor = EXTRACTOR_MAP.get(code)
        if not extractor:
            continue

        docx_path = doc_map[code]
        try:
            doc = Document(str(docx_path))
            extracted = extractor(doc, config)
            for key, value in extracted.items():
                # Only update if we got meaningful data
                if value and value != config.get(key):
                    if verbose:
                        count_info = f"{len(value)} items" if isinstance(value, list) else "updated"
                        if isinstance(value, dict):
                            count_info = "updated"
                        print(f"    [{code:4s}] OK {key}: {count_info}")
                    config[key] = value
                else:
                    if verbose:
                        print(f"    [{code:4s}] -- {key}: no change")
        except Exception as e:
            print(f"    [{code:4s}] ERROR processing {docx_path.name}: {e}")
            if verbose:
                import traceback
                traceback.print_exc()

    # Ensure output_path is always present
    if "output_path" not in config:
        config["output_path"] = str(WORKSPACE_ROOT / "projects")

    return config


# ─── Config I/O ───────────────────────────────────────────────────────────────

def load_existing_config(config_path: Path) -> dict:
    """Load existing config, return empty template structure if not found."""
    if config_path.exists():
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    # Return minimal defaults
    return {
        "project":    {"code": "", "name": "", "organization": ""},
        "team":       {},
        "stakeholders": [],
        "tech_stack": {},
        "milestones": [],
        "requirements": [],
        "design_components": [],
        "database_tables": [],
        "test_cases": [],
        "defects": [],
        "risks": [],
        "change_requests": [],
        "versions": [],
        "meetings": [],
        "training_sessions": [],
        "incidents": [],
        "deployments": [],
        "audit": {},
        "capas": [],
        "sow": {},
        "output_path": str(WORKSPACE_ROOT / "projects"),
    }


def write_config(config: dict, output_path: Path, backup: bool = True):
    """Write config JSON. Creates backup of existing file if requested."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if backup and output_path.exists():
        bak_path = output_path.with_suffix(".json.bak")
        shutil.copy2(str(output_path), str(bak_path))
        print(f"  Backup: {bak_path.name}")
    with open(str(output_path), "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    print(f"  Written: {output_path}")


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="doc_to_config.py — Reverse-extract config from existing .docx files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python tools/doc_to_config.py --project Dvets
  python tools/doc_to_config.py --project HRMS --output configs/HRMS/HRMS_config_new.json
  python tools/doc_to_config.py --project Dvets --no-backup --verbose
  python tools/doc_to_config.py --project GateSurvey --dry-run
        """,
    )
    parser.add_argument("--project",   required=True, help="Project folder name under projects/")
    parser.add_argument("--output",    default=None,  help="Output config path (default: configs/[Project]/[Project]_config.json)")
    parser.add_argument("--no-backup", action="store_true", help="Skip backup of existing config")
    parser.add_argument("--dry-run",   action="store_true", help="Extract and print result, do not write file")
    parser.add_argument("--verbose",   action="store_true", help="Show per-field extraction details")
    args = parser.parse_args()

    project_name = args.project
    project_path = WORKSPACE_ROOT / "projects" / project_name

    if not project_path.exists():
        print(f"ERROR: Project folder not found: {project_path}")
        sys.exit(1)

    # Determine output path
    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = WORKSPACE_ROOT / output_path
    else:
        output_path = WORKSPACE_ROOT / "configs" / project_name / f"{project_name}_config.json"

    # Load existing config as base
    existing_config = load_existing_config(output_path)

    print("=" * 60)
    print(f"  doc_to_config — Extracting config from documents")
    print(f"  Project : {project_name}")
    print(f"  Source  : {project_path}")
    print(f"  Config  : {output_path}")
    print("=" * 60)

    # Extract
    new_config = extract_from_project(project_path, existing_config, verbose=args.verbose)

    if args.dry_run:
        print("\n  [DRY RUN] Extracted config (not written):")
        print(json.dumps(new_config, ensure_ascii=False, indent=2)[:4000])
        print("  (...truncated at 4000 chars for display)")
        return

    # Write
    write_config(new_config, output_path, backup=not args.no_backup)

    # ── Mockup-name validation ─────────────────────────────────────────────────
    try:
        import importlib.util as _ilu
        _vc_path = Path(__file__).parent / "validate_config.py"
        _spec = _ilu.spec_from_file_location("validate_config", _vc_path)
        _vc = _ilu.module_from_spec(_spec)
        _spec.loader.exec_module(_vc)
        _issues = _vc.detect_mock_values(new_config)
        if _issues:
            print()
            _vc.print_report(project_name, _issues)
            print(f"  Tip: python tools/validate_config.py --config \"{output_path}\" --fix")
            print("       to interactively replace placeholder names with real First Last names.")
    except Exception:
        pass  # validation is best-effort; never break the main extraction flow

    # Summary
    print()
    print("  Extraction summary:")
    sections = ["project", "team", "requirements", "test_cases", "defects",
                "risks", "change_requests", "versions", "meetings",
                "training_sessions", "incidents", "audit", "capas"]
    for s in sections:
        val = new_config.get(s)
        if isinstance(val, list):
            print(f"    {s:25s}: {len(val)} items")
        elif isinstance(val, dict) and val:
            print(f"    {s:25s}: updated")
        else:
            print(f"    {s:25s}: (unchanged)")
    print()
    print("  Next: regenerate documents with the updated config:")
    print(f"    python generator/generate_iso_docs.py --config \"{output_path}\"")
    print("=" * 60)


if __name__ == "__main__":
    main()
