"""
audit_snapshot.py — Token-Efficient ISO 29110 Audit Tool
=========================================================
Extracts audit-relevant data from .docx files → compact JSON snapshot (~5-10KB)
instead of loading full 50KB+ documents into AI context.

Usage:
    python tools/audit_snapshot.py --project Dvets
    python tools/audit_snapshot.py --project Dvets --output custom_output.json

Output:
    projects/[ProjectCode]/[ProjectCode]_snapshot.json
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Expected document suffixes per folder
# ---------------------------------------------------------------------------
EXPECTED = {
    "01_Project_Management":   ["PP", "MM"],
    "02_Requirements_Analysis": ["BRD", "RTM"],
    "03_Design_Architecture":  ["SDD", "DBD"],
    "04_Development":          ["CS", "CRR"],
    "05_Testing_QA":           ["TP", "TCR", "BDL"],
    "06_Deployment_Training":  ["UM", "TR"],
    "07_Support_Maintenance":  ["ISL"],
    "08_Change_Logs_Versioning": ["CRF", "VRN"],
    "09_Risk_Management":      ["RR"],
    "10_Regulatory_Compliance": ["IC", "AR", "CAPA"],
}

PLACEHOLDER_PATTERNS = [
    r"\[ระบุ[^\]]*\]",
    r"\(แนบ[^\)]*\)",
    r"\[e\.g\.[^\]]*\]",
    r"_{3,}",
    r"xxx-xxx",
    r"\[TBD\]",
    r"\[TODO\]",
    r"<ระบุ[^>]*>",
    r"\[กรอก[^\]]*\]",
]

ID_PATTERNS = {
    "REQ":  r"REQ-\d+",
    "COMP": r"COMP-\d+",
    "TC":   r"TC-\d+",
    "RISK": r"RISK-\d+",
    "BUG":  r"BUG-\d+",
    "CRF":  r"CRF-\d+",
    "MS":   r"MS-\d+",
    "INC":  r"INC-\d+",
    "ISS":  r"ISS-\d+",
}


def find_placeholders(text: str) -> list[str]:
    """Return list of placeholder strings found in text."""
    found = []
    for pat in PLACEHOLDER_PATTERNS:
        matches = re.findall(pat, text, re.IGNORECASE)
        found.extend(matches)
    return found


def extract_ids(text: str) -> dict[str, list[str]]:
    """Extract all typed IDs (REQ-xxx, BUG-xxx, …) from text."""
    result = {}
    for id_type, pat in ID_PATTERNS.items():
        found = sorted(set(re.findall(pat, text, re.IGNORECASE)))
        if found:
            result[id_type] = found
    return result


def get_signature_status(doc) -> dict:
    """
    Detect empty date/approved/signed cells in signature or document-control tables.
    Returns {"missing_signatures": [...], "total_sig_fields": N}
    """
    missing = []
    total = 0

    for table in doc.tables:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                ct = cell.text.strip().lower()
                # Detect signature / date placeholder cells
                if any(kw in ct for kw in ["วันที่:", "ลงชื่อ:", "approved by:", "date:", "signature:"]):
                    total += 1
                    # Check if the next cell (value cell) is empty or dashes
                    if idx + 1 < len(row.cells):
                        val = row.cells[idx + 1].text.strip()
                        if not val or re.fullmatch(r"[_\-\.]+", val):
                            missing.append(cell.text.strip())

    # Also check paragraph-level signature lines
    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.search(r"(วันที่|Date)\s*:\s*_{2,}", txt, re.IGNORECASE):
            total += 1
            missing.append(txt[:80])

    return {
        "missing_signatures": missing,
        "total_sig_fields": total,
    }


def detect_duplicate_ids(id_list: list[str]) -> list[str]:
    """Return IDs that appear more than once."""
    seen = {}
    for id_val in id_list:
        seen[id_val] = seen.get(id_val, 0) + 1
    return [k for k, v in seen.items() if v > 1]


def get_table_issues(doc) -> list[dict]:
    """
    Detect common table problems:
    - Placeholder cells
    - Empty date/approved columns
    - All-Open status (no resolution)
    - Unticked checklists
    - Duplicate IDs in first column
    - Empty required columns
    """
    issues = []

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # Extract headers from first row
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        header_lower = [h.lower() for h in headers]

        first_col_ids = []
        all_status = []
        empty_date_rows = 0
        placeholder_count = 0
        unchecked_count = 0

        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue  # skip header
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)

            # Collect first-column IDs
            if cells:
                first_col_ids.append(cells[0])

            # Check for placeholders
            ph = find_placeholders(row_text)
            if ph:
                placeholder_count += len(ph)

            # Check status column
            for h_idx, h in enumerate(header_lower):
                if "status" in h or "สถานะ" in h:
                    if h_idx < len(cells):
                        all_status.append(cells[h_idx].lower())

            # Check date columns for empty values
            for h_idx, h in enumerate(header_lower):
                if any(k in h for k in ["date", "วันที่", "due", "กำหนด"]):
                    if h_idx < len(cells) and not cells[h_idx]:
                        empty_date_rows += 1

            # Check unticked checkboxes (☐ or [ ])
            if "☐" in row_text or "[ ]" in row_text:
                unchecked_count += 1

        # Report issues
        if placeholder_count:
            issues.append({
                "table_index": t_idx,
                "issue": "placeholder_cells",
                "count": placeholder_count,
            })

        if all_status and all(s in ["open", "เปิด", "pending", "รอ"] for s in all_status if s):
            issues.append({
                "table_index": t_idx,
                "issue": "all_open_status",
                "note": f"All {len(all_status)} rows have Open/Pending status",
            })

        if empty_date_rows:
            issues.append({
                "table_index": t_idx,
                "issue": "empty_date_columns",
                "count": empty_date_rows,
            })

        if unchecked_count:
            issues.append({
                "table_index": t_idx,
                "issue": "unticked_checkboxes",
                "count": unchecked_count,
            })

        # Check duplicate IDs
        dups = detect_duplicate_ids([i for i in first_col_ids if re.match(r"[A-Z]+-\d+", i)])
        if dups:
            issues.append({
                "table_index": t_idx,
                "issue": "duplicate_ids",
                "ids": dups,
            })

    return issues


def analyze_doc(filepath: Path) -> dict:
    """
    Extract audit-relevant data from a single .docx file.
    Returns a compact summary dict.
    """
    result = {
        "file": filepath.name,
        "path": str(filepath),
        "size_kb": round(filepath.stat().st_size / 1024, 1),
        "headings": [],
        "ids_found": {},
        "placeholders": [],
        "signature_status": {},
        "table_issues": [],
        "version_history": [],
        "has_cover_page": False,
        "has_toc": False,
        "errors": [],
    }

    try:
        doc = Document(str(filepath))
    except Exception as e:
        result["errors"].append(f"Cannot open: {e}")
        return result

    full_text = "\n".join(p.text for p in doc.paragraphs)

    # Extract headings
    result["headings"] = [
        {"level": int(p.style.name[-1]) if p.style.name[-1].isdigit() else 1,
         "text": p.text.strip()}
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]

    # Cover page detection
    result["has_cover_page"] = any(
        kw in full_text[:500].lower()
        for kw in ["document id", "document control", "prepared by", "version history", "cover"]
    )

    # TOC detection
    result["has_toc"] = "table of contents" in full_text.lower() or "สารบัญ" in full_text

    # Extract IDs
    result["ids_found"] = extract_ids(full_text)

    # Placeholders
    all_ph = find_placeholders(full_text)
    unique_ph = list(dict.fromkeys(all_ph))  # deduplicate preserving order
    result["placeholders"] = unique_ph[:20]  # cap at 20

    # Signature status
    result["signature_status"] = get_signature_status(doc)

    # Table issues
    result["table_issues"] = get_table_issues(doc)

    # Version history (look for table with 'version' header)
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if any("version" in h or "เวอร์ชัน" in h for h in headers):
            vers = []
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    vers.append(cells)
            result["version_history"] = vers[:10]
            break

    return result


def check_missing_docs(folder_path: Path) -> dict:
    """
    Check which expected documents are missing from a folder.
    Returns {"folder": name, "found": [...], "missing": [...]}
    """
    folder_name = folder_path.name

    # Find expected codes for this folder
    expected_codes = []
    for prefix, codes in EXPECTED.items():
        if folder_name.startswith(prefix[:2]):  # match by number prefix
            expected_codes = codes
            break
    # Try exact match too
    if not expected_codes and folder_name in EXPECTED:
        expected_codes = EXPECTED[folder_name]

    docx_files = list(folder_path.glob("*.docx"))
    found_codes = []
    for f in docx_files:
        stem = f.stem.upper()
        for code in EXPECTED.get(folder_name, []):
            if code in stem:
                found_codes.append(code)

    missing = [c for c in expected_codes if c not in found_codes]

    return {
        "folder": folder_name,
        "docx_count": len(docx_files),
        "found_codes": found_codes,
        "expected_codes": expected_codes,
        "missing_codes": missing,
    }


def scan_project(project_path: Path) -> dict:
    """
    Scan all folders in a project and build a compact audit snapshot.
    """
    snapshot = {
        "project": project_path.name,
        "snapshot_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "project_path": str(project_path),
        "summary": {
            "total_folders": 0,
            "total_docs": 0,
            "total_placeholders": 0,
            "docs_with_missing_sigs": 0,
            "docs_with_table_issues": 0,
            "missing_docs": [],
            "all_ids": {},
        },
        "folders": [],
    }

    all_ids: dict[str, list[str]] = {}

    # Iterate over expected folders in order
    for folder_name in EXPECTED:
        folder_path = project_path / folder_name
        if not folder_path.exists():
            snapshot["summary"]["missing_docs"].append(f"FOLDER MISSING: {folder_name}")
            continue

        snapshot["summary"]["total_folders"] += 1

        folder_entry = {
            "folder": folder_name,
            "coverage": check_missing_docs(folder_path),
            "documents": [],
        }

        docx_files = sorted(folder_path.glob("*.docx"))
        for docx_file in docx_files:
            snapshot["summary"]["total_docs"] += 1
            doc_data = analyze_doc(docx_file)

            # Accumulate stats
            ph_count = len(doc_data.get("placeholders", []))
            snapshot["summary"]["total_placeholders"] += ph_count

            if doc_data["signature_status"].get("missing_signatures"):
                snapshot["summary"]["docs_with_missing_sigs"] += 1

            if doc_data["table_issues"]:
                snapshot["summary"]["docs_with_table_issues"] += 1

            # Merge IDs
            for id_type, id_list in doc_data.get("ids_found", {}).items():
                if id_type not in all_ids:
                    all_ids[id_type] = []
                all_ids[id_type].extend(id_list)

            # Keep only essential fields for the snapshot (reduce size)
            compact = {
                "file": doc_data["file"],
                "size_kb": doc_data["size_kb"],
                "headings": [h["text"] for h in doc_data["headings"][:10]],
                "ids_found": doc_data["ids_found"],
                "placeholder_count": ph_count,
                "placeholders_sample": doc_data["placeholders"][:5],
                "missing_signatures": doc_data["signature_status"].get("missing_signatures", []),
                "table_issues": doc_data["table_issues"],
                "version_history_rows": len(doc_data.get("version_history", [])),
                "has_cover_page": doc_data["has_cover_page"],
            }
            if doc_data.get("errors"):
                compact["errors"] = doc_data["errors"]

            folder_entry["documents"].append(compact)

        snapshot["folders"].append(folder_entry)

    # Deduplicate and sort all IDs
    snapshot["summary"]["all_ids"] = {
        k: sorted(set(v)) for k, v in all_ids.items()
    }

    return snapshot


def generate_text_report(snapshot: dict) -> str:
    """Generate a human-readable summary from the snapshot (for quick review)."""
    lines = []
    lines.append(f"{'='*60}")
    lines.append(f"AUDIT SNAPSHOT REPORT")
    lines.append(f"Project : {snapshot['project']}")
    lines.append(f"Date    : {snapshot['snapshot_date']}")
    lines.append(f"{'='*60}")

    s = snapshot["summary"]
    lines.append(f"\nSUMMARY")
    lines.append(f"  Folders scanned : {s['total_folders']}")
    lines.append(f"  Documents found : {s['total_docs']}")
    lines.append(f"  Total placeholders : {s['total_placeholders']}")
    lines.append(f"  Docs missing sigs  : {s['docs_with_missing_sigs']}")
    lines.append(f"  Docs with table issues : {s['docs_with_table_issues']}")

    if s["missing_docs"]:
        lines.append(f"\nMISSING FOLDERS/DOCS:")
        for m in s["missing_docs"]:
            lines.append(f"  ⚠ {m}")

    lines.append(f"\nALL IDs FOUND:")
    for id_type, ids in s["all_ids"].items():
        lines.append(f"  {id_type}: {', '.join(ids)}")

    lines.append(f"\nFOLDER DETAILS:")
    for folder in snapshot["folders"]:
        cov = folder["coverage"]
        missing = cov.get("missing_codes", [])
        status = "✅" if not missing else f"⚠️  Missing: {missing}"
        lines.append(f"\n  📁 {folder['folder']} — {cov['docx_count']} docs {status}")

        for doc in folder["documents"]:
            ph = doc["placeholder_count"]
            sigs = len(doc.get("missing_signatures", []))
            tbl = len(doc.get("table_issues", []))
            flags = []
            if ph:
                flags.append(f"{ph} placeholders")
            if sigs:
                flags.append(f"{sigs} missing sigs")
            if tbl:
                flags.append(f"{tbl} table issues")
            flag_str = " | ".join(flags) if flags else "clean"
            lines.append(f"     📄 {doc['file']} ({doc['size_kb']}KB) — {flag_str}")

    lines.append(f"\n{'='*60}")
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Extract audit-relevant data from ISO project .docx files → compact JSON"
    )
    parser.add_argument(
        "--project", required=True,
        help="Project code (e.g. Dvets). Folder must exist under projects/"
    )
    parser.add_argument(
        "--output", default=None,
        help="Custom output path for JSON file. Defaults to projects/[Project]/[Project]_snapshot.json"
    )
    parser.add_argument(
        "--report", action="store_true",
        help="Also print a human-readable summary report"
    )
    args = parser.parse_args()

    # Resolve project path
    script_dir = Path(__file__).resolve().parent
    workspace_root = script_dir.parent
    project_path = workspace_root / "projects" / args.project

    if not project_path.exists():
        print(f"ERROR: Project folder not found: {project_path}")
        sys.exit(1)

    print(f"Scanning project: {args.project}")
    print(f"Path: {project_path}")

    snapshot = scan_project(project_path)

    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_dir = workspace_root / "configs" / args.project
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{args.project}_snapshot.json"

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)

    size_kb = round(output_path.stat().st_size / 1024, 1)
    print(f"\n✅ Snapshot saved: {output_path}")
    print(f"   Size: {size_kb} KB  (vs ~50KB+ for raw .docx content)")
    print(f"   Documents: {snapshot['summary']['total_docs']}")
    print(f"   Placeholders: {snapshot['summary']['total_placeholders']}")
    print(f"   Docs with missing signatures: {snapshot['summary']['docs_with_missing_sigs']}")

    if args.report:
        print()
        print(generate_text_report(snapshot))


if __name__ == "__main__":
    main()
