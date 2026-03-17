"""
doc_patcher.py — ISO 29110 Document Auto-Patcher
==================================================
Applies AI-generated patches to .docx files automatically.
Always creates a .backup before modifying any file.

Usage:
    # Dry run (preview changes, no file modifications)
    python tools/doc_patcher.py --project Dvets --patches configs/Dvets/Dvets_patches.json --dry-run

    # Apply patches
    python tools/doc_patcher.py --project Dvets --patches configs/Dvets/Dvets_patches.json

    # Apply and skip backup (not recommended)
    python tools/doc_patcher.py --project Dvets --patches configs/Dvets/Dvets_patches.json --no-backup

Patch JSON format (Dvets_patches.json):
    {
      "project": "Dvets",
      "patches": [
        {
          "file_pattern": "01_Project_Management/DVETS-01-PP-v1.0.docx",
          "action": "replace_in_paragraph",
          "find": "[ระบุวันที่อนุมัติ]",
          "replace": "2025-10-01",
          "condition_paragraph_contains": "วันที่อนุมัติ"
        },
        {
          "file_pattern": "05_Testing_QA/DVETS-05-TCR-v1.0.docx",
          "action": "fill_table_cell",
          "row_contains": "TC-001",
          "col_header": "Result",
          "value": "Pass"
        },
        {
          "file_pattern": "07_Support_Maintenance/DVETS-07-ISL-v1.0.docx",
          "action": "fix_duplicate_id",
          "id_prefix": "ISS",
          "col_index": 0
        }
      ]
    }

Supported actions:
    replace_in_paragraph   — find/replace text in paragraphs
    fill_table_cell        — fill an empty cell identified by row keyword + column header
    update_table_status    — change a status value in a specific row
    fill_empty_dates       — fill blank date fields matching a pattern
    fix_duplicate_id       — renumber duplicate IDs in a column
    append_table_row       — append a new row to a table identified by its headers
"""

import argparse
import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Action Handlers
# ---------------------------------------------------------------------------

def action_replace_in_paragraph(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Find text in paragraphs and replace it.
    Optional: only replace if paragraph also contains condition_paragraph_contains.
    Returns list of changes made.
    """
    find_text = patch["find"]
    replace_text = patch["replace"]
    condition = patch.get("condition_paragraph_contains", "")
    changes = []

    for para in doc.paragraphs:
        if find_text in para.text:
            if condition and condition not in para.text:
                continue
            if dry_run:
                changes.append(f"[DRY] Replace '{find_text}' → '{replace_text}' in paragraph: {para.text[:80]}")
            else:
                # Preserve runs by replacing in each run
                for run in para.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                changes.append(f"Replaced '{find_text}' → '{replace_text}' in paragraph: {para.text[:80]}")

    # Also search in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if find_text in para.text:
                        if condition and condition not in para.text:
                            continue
                        if dry_run:
                            changes.append(f"[DRY] Replace '{find_text}' → '{replace_text}' in table cell: {para.text[:80]}")
                        else:
                            for run in para.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                            changes.append(f"Replaced '{find_text}' → '{replace_text}' in table cell")

    return changes


def action_fill_table_cell(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Find a table row containing row_contains text, then fill the cell
    under col_header with value (only if the cell is currently empty).
    """
    row_contains = patch["row_contains"]
    col_header = patch["col_header"].lower()
    value = patch["value"]
    overwrite_empty_only = patch.get("overwrite_empty_only", True)
    changes = []

    for table in doc.tables:
        if not table.rows:
            continue
        # Find header row index and column index
        header_row = table.rows[0]
        col_idx = None
        for i, cell in enumerate(header_row.cells):
            if col_header in cell.text.lower():
                col_idx = i
                break
        if col_idx is None:
            continue

        # Find target data rows
        for r_idx, row in enumerate(table.rows[1:], start=1):
            row_text = " ".join(c.text for c in row.cells)
            if row_contains.lower() in row_text.lower():
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    current = cell.text.strip()
                    if overwrite_empty_only and current:
                        changes.append(f"SKIP (not empty, row {r_idx}, col '{col_header}'): '{current}'")
                        continue
                    if dry_run:
                        changes.append(f"[DRY] Fill cell row {r_idx} col '{col_header}' with '{value}' (was: '{current}')")
                    else:
                        # Clear and set text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = ""
                        if cell.paragraphs:
                            cell.paragraphs[0].text = value
                        else:
                            cell.add_paragraph(value)
                        changes.append(f"Filled cell row {r_idx} col '{col_header}' with '{value}'")

    return changes


def action_update_table_status(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Find rows in a table where row_contains matches, and change
    the cell in col_header from old_value to new_value.
    """
    row_contains = patch.get("row_contains", "")
    col_header = patch["col_header"].lower()
    old_value = patch.get("old_value", "")
    new_value = patch["new_value"]
    changes = []

    for table in doc.tables:
        if not table.rows:
            continue
        header_row = table.rows[0]
        col_idx = None
        for i, cell in enumerate(header_row.cells):
            if col_header in cell.text.lower():
                col_idx = i
                break
        if col_idx is None:
            continue

        for r_idx, row in enumerate(table.rows[1:], start=1):
            row_text = " ".join(c.text for c in row.cells)
            if row_contains and row_contains.lower() not in row_text.lower():
                continue
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                current = cell.text.strip()
                if old_value and current.lower() != old_value.lower():
                    continue
                if dry_run:
                    changes.append(f"[DRY] Update status row {r_idx}: '{current}' → '{new_value}'")
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    if cell.paragraphs:
                        cell.paragraphs[0].text = new_value
                    else:
                        cell.add_paragraph(new_value)
                    changes.append(f"Updated status row {r_idx}: '{current}' → '{new_value}'")

    return changes


def action_fill_empty_dates(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Find paragraph patterns like 'วันที่: ___' or 'Date: ___' and fill them.
    date_pattern: regex to match the line
    date_value: the value to insert
    """
    date_pattern = patch.get("date_pattern", r"(วันที่|Date)\s*:\s*_{2,}")
    date_value = patch.get("date_value", datetime.now().strftime("%Y-%m-%d"))
    changes = []

    def fill_paragraph(para):
        if re.search(date_pattern, para.text, re.IGNORECASE):
            new_text = re.sub(
                r"_{2,}",
                date_value,
                para.text,
                flags=re.IGNORECASE,
            )
            if dry_run:
                changes.append(f"[DRY] Fill date in: '{para.text[:80]}' → '{new_text[:80]}'")
            else:
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)
                changes.append(f"Filled date: '{new_text[:80]}'")

    for para in doc.paragraphs:
        fill_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fill_paragraph(para)

    return changes


def action_fix_duplicate_id(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Find duplicate IDs in a specified column (by index or header name)
    and renumber them sequentially.
    id_prefix: e.g. "ISS", "REQ", "BUG"
    col_index: 0-based column index (optional, default 0)
    col_header: header name to find column (alternative to col_index)
    """
    id_prefix = patch["id_prefix"]
    col_index = patch.get("col_index", None)
    col_header = patch.get("col_header", "").lower()
    changes = []

    id_pattern = re.compile(rf"^{re.escape(id_prefix)}-(\d+)$", re.IGNORECASE)

    for table in doc.tables:
        if not table.rows:
            continue

        # Determine column index
        target_col = col_index
        if target_col is None and col_header:
            for i, cell in enumerate(table.rows[0].cells):
                if col_header in cell.text.lower():
                    target_col = i
                    break
        if target_col is None:
            target_col = 0

        # Collect existing IDs
        seen_ids = {}
        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue
            if target_col < len(row.cells):
                cell_text = row.cells[target_col].text.strip()
                m = id_pattern.match(cell_text)
                if m:
                    num = int(m.group(1))
                    seen_ids.setdefault(num, []).append(r_idx)

        # Renumber duplicates
        for num, row_idxs in seen_ids.items():
            if len(row_idxs) <= 1:
                continue
            # Keep first, renumber subsequent
            max_num = max(seen_ids.keys())
            for extra_r_idx in row_idxs[1:]:
                max_num += 1
                new_id = f"{id_prefix.upper()}-{max_num:03d}"
                old_id = f"{id_prefix.upper()}-{num:03d}"
                cell = table.rows[extra_r_idx].cells[target_col]
                if dry_run:
                    changes.append(f"[DRY] Renumber duplicate {old_id} → {new_id} at row {extra_r_idx}")
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(old_id, new_id).replace(
                                f"{id_prefix.lower()}-{num:03d}", new_id
                            )
                    changes.append(f"Renumbered {old_id} → {new_id} at row {extra_r_idx}")

    return changes


def action_append_table_row(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Append a new row to a table identified by its headers.
    table_headers: list of strings that must all appear in the first row
    row_data: list of cell values to insert
    """
    table_headers = [h.lower() for h in patch.get("table_headers", [])]
    row_data = patch.get("row_data", [])
    changes = []

    if not table_headers or not row_data:
        return ["ERROR: append_table_row requires 'table_headers' and 'row_data'"]

    for table in doc.tables:
        if not table.rows:
            continue
        first_row_text = " ".join(c.text.lower() for c in table.rows[0].cells)
        if all(h in first_row_text for h in table_headers):
            if dry_run:
                changes.append(f"[DRY] Append row: {row_data}")
            else:
                new_row = table.add_row()
                for i, value in enumerate(row_data):
                    if i < len(new_row.cells):
                        new_row.cells[i].text = str(value)
                changes.append(f"Appended row: {row_data}")
            break

    return changes


def action_replace_checkbox(doc, patch: dict, dry_run: bool) -> list[str]:
    """
    Replace checkbox characters in table cells.
    find: text to find (default: "☐")
    replace_with: text to replace with (default: "☑")
    table_index: only apply to this specific table (0-based), optional
    col_index: only in this column (0-based), optional
    col_header: find column by header text, optional
    condition_row_contains: only rows containing this text, 'ALL' = all data rows
    """
    find_char = patch.get("find", "☐")
    replace_char = patch.get("replace_with", "☑")
    target_table_index = patch.get("table_index", None)
    col_index_param = patch.get("col_index", None)
    col_header_param = patch.get("col_header", "").lower()
    condition_row = patch.get("condition_row_contains", "ALL")
    changes = []

    for t_idx, table in enumerate(doc.tables):
        if target_table_index is not None and t_idx != target_table_index:
            continue
        if not table.rows:
            continue

        # Determine column index from header
        col_idx = col_index_param
        if col_idx is None and col_header_param:
            for i, cell in enumerate(table.rows[0].cells):
                if col_header_param in cell.text.lower():
                    col_idx = i
                    break

        for r_idx, row in enumerate(table.rows):
            if col_header_param and r_idx == 0:
                continue  # skip header row

            if condition_row != "ALL":
                row_text = " ".join(c.text for c in row.cells)
                if condition_row.lower() not in row_text.lower():
                    continue

            cells_to_process = ([row.cells[col_idx]] if col_idx is not None and col_idx < len(row.cells)
                                 else list(row.cells))

            for cell in cells_to_process:
                for para in cell.paragraphs:
                    if find_char in para.text:
                        new_text = para.text.replace(find_char, replace_char, 1)
                        if dry_run:
                            changes.append(f"[DRY] Table {t_idx} row {r_idx}: '{para.text[:60]}' → '{new_text[:60]}'")
                        else:
                            full_new = para.text.replace(find_char, replace_char, 1)
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = full_new
                            else:
                                para.add_run(full_new)
                            changes.append(f"Table {t_idx} row {r_idx}: '{para.text[:40]}' → checkbox updated")
    return changes


# ---------------------------------------------------------------------------
# Action dispatcher
# ---------------------------------------------------------------------------

ACTION_MAP = {
    "replace_in_paragraph":  action_replace_in_paragraph,
    "fill_table_cell":       action_fill_table_cell,
    "update_table_status":   action_update_table_status,
    "fill_empty_dates":      action_fill_empty_dates,
    "fix_duplicate_id":      action_fix_duplicate_id,
    "append_table_row":      action_append_table_row,
    "replace_checkbox":      action_replace_checkbox,
}


# ---------------------------------------------------------------------------
# Main patch application logic
# ---------------------------------------------------------------------------

def resolve_docx_path(project_path: Path, file_pattern: str) -> Path | None:
    """
    Resolve a file_pattern to an absolute path.
    Supports:
    - Relative path: "01_Project_Management/DVETS-01-PP-v1.0.docx"
    - Glob pattern: "01_Project_Management/*.docx"
    - Filename only: "DVETS-01-PP-v1.0.docx" (searched recursively)
    """
    target = project_path / file_pattern
    if target.exists():
        return target

    # Try glob
    results = list(project_path.glob(file_pattern))
    if results:
        return results[0]

    # Try recursive search by filename
    filename = Path(file_pattern).name
    results = list(project_path.rglob(filename))
    if results:
        return results[0]

    return None


def apply_patches(
    project_path: Path,
    patches: list[dict],
    dry_run: bool = False,
    no_backup: bool = False,
) -> dict:
    """
    Apply a list of patches to .docx files in the project.
    Groups patches by file_pattern for efficiency (one save per file).
    Returns a results summary dict.
    """
    results = {
        "dry_run": dry_run,
        "project": project_path.name,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "files_processed": 0,
        "files_backed_up": 0,
        "total_changes": 0,
        "errors": [],
        "file_results": [],
    }

    # Group patches by file
    by_file: dict[str, list[dict]] = {}
    for patch in patches:
        fp = patch.get("file_pattern", "")
        if not fp:
            results["errors"].append("Patch missing 'file_pattern'")
            continue
        by_file.setdefault(fp, []).append(patch)

    for file_pattern, file_patches in by_file.items():
        docx_path = resolve_docx_path(project_path, file_pattern)
        if not docx_path:
            results["errors"].append(f"File not found: {file_pattern}")
            continue

        file_result = {
            "file": str(docx_path.relative_to(project_path)),
            "changes": [],
            "errors": [],
        }

        try:
            doc = Document(str(docx_path))
        except Exception as e:
            file_result["errors"].append(f"Cannot open: {e}")
            results["file_results"].append(file_result)
            continue

        # Apply each patch
        doc_modified = False
        for patch in file_patches:
            action_name = patch.get("action", "")
            handler = ACTION_MAP.get(action_name)
            if not handler:
                file_result["errors"].append(f"Unknown action: '{action_name}'")
                continue
            try:
                changes = handler(doc, patch, dry_run)
                file_result["changes"].extend(changes)
                if changes and any("[DRY]" not in c for c in changes):
                    doc_modified = True
            except Exception as e:
                file_result["errors"].append(f"Action '{action_name}' failed: {e}")

        # Save if modified (not dry run)
        if doc_modified and not dry_run:
            # Create backup
            if not no_backup:
                backup_path = docx_path.with_suffix(".backup.docx")
                shutil.copy2(str(docx_path), str(backup_path))
                results["files_backed_up"] += 1
                file_result["backup"] = str(backup_path.relative_to(project_path))

            try:
                doc.save(str(docx_path))
                file_result["saved"] = True
            except Exception as e:
                file_result["errors"].append(f"Cannot save: {e}")
                file_result["saved"] = False

        results["files_processed"] += 1
        results["total_changes"] += len([c for c in file_result["changes"] if "SKIP" not in c])
        results["file_results"].append(file_result)

    return results


def print_results(results: dict):
    """Print results in human-readable format."""
    mode = "DRY RUN" if results["dry_run"] else "APPLIED"
    print(f"\n{'='*60}")
    print(f"DOC PATCHER — {mode}")
    print(f"Project  : {results['project']}")
    print(f"Timestamp: {results['timestamp']}")
    print(f"{'='*60}")
    print(f"Files processed : {results['files_processed']}")
    print(f"Files backed up : {results['files_backed_up']}")
    print(f"Total changes   : {results['total_changes']}")

    if results["errors"]:
        print(f"\nGLOBAL ERRORS:")
        for e in results["errors"]:
            print(f"  ❌ {e}")

    for fr in results["file_results"]:
        print(f"\n  📄 {fr['file']}")
        if fr.get("backup"):
            print(f"     💾 Backup: {fr['backup']}")
        for ch in fr["changes"]:
            prefix = "  [DRY]" if "[DRY]" in ch else "  ✅"
            print(f"     {prefix} {ch.replace('[DRY] ', '')}")
        for e in fr.get("errors", []):
            print(f"     ❌ {e}")

    print(f"\n{'='*60}")
    if results["dry_run"]:
        print("No files were modified. Run without --dry-run to apply changes.")
    else:
        print("Patches applied. Original files backed up as .backup.docx")


def main():
    parser = argparse.ArgumentParser(
        description="Apply AI-generated patches to ISO .docx files"
    )
    parser.add_argument(
        "--project", required=True,
        help="Project code (e.g. Dvets)"
    )
    parser.add_argument(
        "--patches", required=True,
        help="Path to patches JSON file (absolute or relative to workspace root)"
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Preview changes without modifying any files"
    )
    parser.add_argument(
        "--no-backup", action="store_true",
        help="Skip creating .backup.docx files (not recommended)"
    )
    parser.add_argument(
        "--output-report", default=None,
        help="Save results JSON to this path"
    )
    args = parser.parse_args()

    # Resolve project path
    script_dir = Path(__file__).resolve().parent
    workspace_root = script_dir.parent
    project_path = workspace_root / "projects" / args.project

    if not project_path.exists():
        print(f"ERROR: Project folder not found: {project_path}")
        sys.exit(1)

    # Load patches
    patches_path = Path(args.patches)
    if not patches_path.is_absolute():
        # Try relative to workspace root
        patches_path = workspace_root / args.patches
        if not patches_path.exists():
            # Try relative to project folder
            patches_path = project_path / args.patches
    if not patches_path.exists():
        print(f"ERROR: Patches file not found: {args.patches}")
        sys.exit(1)

    with open(patches_path, encoding="utf-8-sig") as f:
        patch_data = json.load(f)

    patches_list = patch_data.get("patches", patch_data if isinstance(patch_data, list) else [])

    print(f"Loaded {len(patches_list)} patches from: {patches_path}")
    print(f"Project path: {project_path}")
    if args.dry_run:
        print("Mode: DRY RUN (no files will be modified)")

    results = apply_patches(
        project_path=project_path,
        patches=patches_list,
        dry_run=args.dry_run,
        no_backup=args.no_backup,
    )

    print_results(results)

    if args.output_report:
        report_path = Path(args.output_report)
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"\nResults saved to: {report_path}")


if __name__ == "__main__":
    main()
